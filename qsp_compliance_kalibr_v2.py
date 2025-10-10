"""
QSP Compliance Checker - Kalibr SDK v2 Implementation
Multi-Model AI Integration for ISO 13485:2024 Compliance Analysis

This replaces our 700+ line custom MCP implementation with clean Kalibr SDK functions
that automatically work with GPT, Claude, Gemini, and Copilot.
"""

from kalibr import KalibrApp
from kalibr.types import FileUpload, Session, StreamingResponse, WorkflowState
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime, timezone
import uuid
import json
import base64
import io
import re
from docx import Document
import asyncio

# Import existing LLM integration (reuse our existing AI analysis)
try:
    from emergentintegrations.llm.chat import LlmChat, UserMessage
except ImportError:
    print("Note: emergentintegrations not available. AI analysis will be mocked.")
    LlmChat = None

# Initialize Enhanced Kalibr App
app = KalibrApp(
    title="QSP Compliance Checker",
    base_url="http://localhost:8000"
)

# Data Models
class QSPDocument(BaseModel):
    id: str
    filename: str
    content: str
    sections: Dict[str, str]
    upload_date: str
    processed: bool = True

class ISOSummary(BaseModel):
    id: str
    framework: str = "ISO_13485"
    version: str = "2024_summary"
    content: str
    new_clauses: List[Dict[str, str]]
    modified_clauses: List[Dict[str, str]]
    upload_date: str

class ClauseMapping(BaseModel):
    id: str
    qsp_id: str
    qsp_filename: str
    section_title: str
    iso_clause: str
    confidence_score: float
    evidence_text: str

class ComplianceGap(BaseModel):
    id: str
    qsp_id: str
    qsp_filename: str
    iso_clause: str
    gap_type: str  # "missing", "partial", "outdated"
    description: str
    severity: str  # "high", "medium", "low"
    recommendations: List[str]

# In-memory storage (replace with database in production)
qsp_documents: Dict[str, QSPDocument] = {}
iso_summaries: Dict[str, ISOSummary] = {}
clause_mappings: Dict[str, ClauseMapping] = {}
compliance_gaps: Dict[str, ComplianceGap] = {}
analysis_history: List[Dict] = []

# Helper functions
def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX file: {e}")

def parse_document_sections(content: str, filename: str) -> Dict[str, str]:
    """Parse document content into logical sections"""
    sections = {}
    
    # Split by common section patterns
    patterns = [
        r'\b\d+\.\s+[A-Z][^.]*\n',  # Numbered sections
        r'\b[A-Z][^.]{10,50}:?\n',   # Header patterns
        r'^[A-Z\s]{3,}\n',           # All caps headers
    ]
    
    current_section = "Introduction"
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line is a section header
        is_header = False
        for pattern in patterns:
            if re.match(pattern, line + '\n'):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = line.replace(':', '').strip()
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(line)
    
    # Save final section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    # If no sections found, treat entire document as one section
    if not sections:
        sections[f"{filename} - Full Content"] = content
    
    return sections

async def analyze_with_ai(qsp_content: str, qsp_filename: str, iso_clauses: List[str]) -> List[Dict]:
    """Use AI to analyze QSP content against ISO clauses"""
    if not LlmChat:
        # Mock AI analysis for demo
        return [
            {
                "iso_clause": iso_clauses[0] if iso_clauses else "4.1 General requirements",
                "confidence_score": 0.85,
                "evidence_text": qsp_content[:200] + "...",
                "explanation": "Document contains relevant quality management requirements"
            }
        ]
    
    try:
        # Use existing AI integration
        chat = LlmChat(
            session_id=f"qsp_analysis_{uuid.uuid4()}",
            system_message="""You are an expert in ISO 13485 compliance for medical devices. 
            Analyze QSP document sections and map them to relevant ISO 13485:2024 clauses.
            
            Return results as a JSON array with this structure:
            [
                {
                    "iso_clause": "4.1 General requirements",
                    "confidence_score": 0.85,
                    "evidence_text": "relevant excerpt from QSP",
                    "explanation": "brief explanation of why this maps"
                }
            ]"""
        ).with_model("openai", "gpt-4o")
        
        iso_clauses_text = "\n".join([f"- {clause}" for clause in iso_clauses])
        
        user_message = UserMessage(
            text=f"""Analyze this QSP document content and map it to relevant ISO 13485:2024 clauses:

QSP Document: {qsp_filename}
Content:
{qsp_content[:2000]}...

ISO 13485:2024 Key Clauses:
{iso_clauses_text}

Provide detailed mappings with confidence scores."""
        )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            clean_response = response.strip()
            if clean_response.startswith('```json'):
                clean_response = clean_response[7:]
            if clean_response.endswith('```'):
                clean_response = clean_response[:-3]
            clean_response = clean_response.strip()
            
            mappings = json.loads(clean_response)
            return mappings if isinstance(mappings, list) else []
        except json.JSONDecodeError:
            return []
            
    except Exception as e:
        print(f"AI analysis error: {e}")
        return []

# ===================  KALIBR ACTIONS ===================

@app.file_handler("upload_qsp_document", [".docx", ".txt"])
async def upload_qsp_document(file: FileUpload):
    """Upload and process a QSP document for ISO 13485 compliance analysis"""
    try:
        # Extract text based on file type
        if file.filename.endswith('.docx'):
            text_content = extract_text_from_docx(file.content)
        else:  # .txt
            text_content = file.content.decode('utf-8')
        
        # Parse into sections
        sections = parse_document_sections(text_content, file.filename)
        
        # Create QSP document
        qsp_doc = QSPDocument(
            id=str(uuid.uuid4()),
            filename=file.filename,
            content=text_content,
            sections=sections,
            upload_date=datetime.now(timezone.utc).isoformat()
        )
        
        # Store in memory
        qsp_documents[qsp_doc.id] = qsp_doc
        
        return {
            "success": True,
            "upload_id": file.upload_id,
            "document_id": qsp_doc.id,
            "filename": file.filename,
            "sections_count": len(sections),
            "content_length": len(text_content),
            "message": f"QSP document '{file.filename}' uploaded and processed successfully. Found {len(sections)} sections ready for analysis."
        }
        
    except Exception as e:
        return {
            "success": False,
            "upload_id": file.upload_id,
            "error": str(e),
            "message": "Failed to process QSP document"
        }

@app.file_handler("upload_iso_summary", [".docx", ".txt", ".pdf"])
async def upload_iso_summary(file: FileUpload):
    """Upload ISO 13485:2024 Summary of Changes for compliance benchmarking"""
    try:
        # Extract text
        if file.filename.endswith('.docx'):
            text_content = extract_text_from_docx(file.content)
        else:
            text_content = file.content.decode('utf-8')
        
        # Parse ISO summary content
        new_clauses = []
        modified_clauses = []
        
        lines = text_content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if 'NEW CLAUSES' in line.upper() or 'NEW REQUIREMENTS' in line.upper():
                current_section = "new"
            elif 'MODIFIED CLAUSES' in line.upper() or 'CHANGED CLAUSES' in line.upper():
                current_section = "modified"
            elif line and current_section and re.match(r'\d+\.\d+', line):
                clause_info = {"clause": line, "description": line}
                if current_section == "new":
                    new_clauses.append(clause_info)
                elif current_section == "modified":
                    modified_clauses.append(clause_info)
        
        # Create ISO summary
        iso_summary = ISOSummary(
            id=str(uuid.uuid4()),
            content=text_content,
            new_clauses=new_clauses,
            modified_clauses=modified_clauses,
            upload_date=datetime.now(timezone.utc).isoformat()
        )
        
        # Store in memory
        iso_summaries[iso_summary.id] = iso_summary
        
        return {
            "success": True,
            "upload_id": file.upload_id,
            "summary_id": iso_summary.id,
            "filename": file.filename,
            "new_clauses_count": len(new_clauses),
            "modified_clauses_count": len(modified_clauses),
            "message": f"ISO summary '{file.filename}' uploaded successfully. Found {len(new_clauses)} new and {len(modified_clauses)} modified clauses."
        }
        
    except Exception as e:
        return {
            "success": False,
            "upload_id": file.upload_id,
            "error": str(e),
            "message": "Failed to process ISO summary document"
        }

@app.action("list_documents", "Get all uploaded QSP documents with their processing status")
def list_documents(format_type: str = "summary", include_content: bool = False):
    """List all QSP documents"""
    try:
        docs = list(qsp_documents.values())
        
        if format_type == "detailed":
            result = []
            for doc in docs:
                doc_info = doc.dict()
                if not include_content:
                    doc_info.pop('content', None)  # Remove large content unless requested
                result.append(doc_info)
            
            return {
                "documents": result,
                "count": len(docs),
                "message": f"Retrieved {len(docs)} documents with detailed information"
            }
        
        # Summary format
        summary = []
        for doc in docs:
            summary.append({
                "id": doc.id,
                "filename": doc.filename,
                "sections": len(doc.sections),
                "upload_date": doc.upload_date,
                "processed": doc.processed,
                "content_preview": doc.content[:100] + "..." if len(doc.content) > 100 else doc.content
            })
        
        return {
            "documents": summary,
            "count": len(docs),
            "iso_summaries_available": len(iso_summaries),
            "message": f"Found {len(docs)} QSP documents. Ready for analysis." if docs else "No documents uploaded yet."
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.stream_action("run_clause_mapping", "Execute AI-powered mapping between QSP sections and ISO 13485 clauses with real-time progress")
async def run_clause_mapping(use_ai: bool = True, min_confidence: float = 0.3):
    """Run AI-powered clause mapping analysis with streaming progress"""
    try:
        if not qsp_documents:
            yield {"error": "No QSP documents found. Upload documents first.", "progress": 0}
            return
        
        # Standard ISO 13485 clauses for mapping
        iso_clauses = [
            "4.1 General requirements",
            "4.2 Documentation requirements",
            "5.1 Management commitment",
            "5.2 Customer focus",
            "5.3 Quality policy",
            "5.4 Planning",
            "5.5 Responsibility, authority and communication",
            "5.6 Management review",
            "6.1 Provision of resources",
            "6.2 Human resources",
            "6.3 Infrastructure",
            "6.4 Work environment",
            "7.1 Planning of product realization",
            "7.2 Customer-related processes",
            "7.3 Design and development",
            "7.4 Purchasing",
            "7.5 Production and service provision",
            "7.6 Control of monitoring and measuring equipment",
            "8.1 General",
            "8.2 Monitoring and measurement",
            "8.3 Control of nonconforming product",
            "8.4 Analysis of data",
            "8.5 Improvement"
        ]
        
        # Clear existing mappings
        clause_mappings.clear()
        
        total_docs = len(qsp_documents)
        total_mappings = 0
        processed_docs = 0
        
        yield {
            "status": "starting",
            "progress": 5,
            "message": f"Starting analysis of {total_docs} documents...",
            "total_documents": total_docs
        }
        
        # Process each QSP document
        for doc in qsp_documents.values():
            processed_docs += 1
            progress = (processed_docs / total_docs) * 80 + 10  # 10-90% range
            
            yield {
                "status": "processing",
                "progress": progress,
                "message": f"Analyzing document: {doc.filename}",
                "current_document": doc.filename,
                "processed_documents": processed_docs
            }
            
            # Process each section
            sections_processed = 0
            total_sections = len(doc.sections)
            
            for section_title, section_content in doc.sections.items():
                if len(section_content) < 50:  # Skip very short sections
                    continue
                
                sections_processed += 1
                
                if use_ai and LlmChat:
                    # AI-powered analysis
                    mappings = await analyze_with_ai(section_content, doc.filename, iso_clauses)
                else:
                    # Fallback simple analysis
                    mappings = []
                    for iso_clause in iso_clauses:
                        confidence = calculate_simple_confidence(section_content, iso_clause)
                        if confidence > min_confidence:
                            mappings.append({
                                "iso_clause": iso_clause,
                                "confidence_score": confidence,
                                "evidence_text": section_content[:200] + "...",
                                "explanation": f"Keyword-based match for {iso_clause}"
                            })
                
                # Store mappings
                for mapping in mappings:
                    if mapping.get('confidence_score', 0) > min_confidence:
                        clause_mapping = ClauseMapping(
                            id=str(uuid.uuid4()),
                            qsp_id=doc.id,
                            qsp_filename=doc.filename,
                            section_title=section_title,
                            iso_clause=mapping['iso_clause'],
                            confidence_score=mapping['confidence_score'],
                            evidence_text=mapping['evidence_text']
                        )
                        
                        clause_mappings[clause_mapping.id] = clause_mapping
                        total_mappings += 1
                
                # Small delay for streaming effect
                await asyncio.sleep(0.1)
        
        # Final results
        yield {
            "status": "completed",
            "progress": 100,
            "message": "Clause mapping analysis completed successfully!",
            "results": {
                "documents_processed": processed_docs,
                "mappings_generated": total_mappings,
                "average_mappings_per_doc": round(total_mappings / processed_docs, 1) if processed_docs > 0 else 0,
                "analysis_method": "AI-powered" if use_ai else "Keyword-based"
            }
        }
        
    except Exception as e:
        yield {
            "status": "error",
            "progress": 0,
            "error": str(e),
            "message": "Analysis failed"
        }

def calculate_simple_confidence(content: str, iso_clause: str) -> float:
    """Simple confidence calculation for fallback analysis"""
    content_lower = content.lower()
    clause_lower = iso_clause.lower()
    
    # Keyword mapping for different clauses
    keywords = {
        "4.1 general": ["requirements", "general", "quality", "management", "system"],
        "4.2 documentation": ["documentation", "document", "records", "procedure", "control"],
        "5.1 management": ["management", "commitment", "leadership", "responsibility", "policy"],
        "5.2 customer": ["customer", "client", "satisfaction", "requirements", "focus"],
        "5.3 quality policy": ["policy", "quality", "objectives", "commitment"],
        "7.3 design": ["design", "development", "product", "specification", "validation"],
        "7.4 purchasing": ["purchasing", "supplier", "vendor", "procurement", "evaluation"],
        "8.1 general": ["monitoring", "measurement", "control", "processes"],
        "8.5 improvement": ["improvement", "corrective", "preventive", "action", "nonconformity"]
    }
    
    for key_phrase, terms in keywords.items():
        if key_phrase in clause_lower:
            matches = sum(1 for term in terms if term in content_lower)
            return min(matches * 0.15, 0.9)  # Max confidence of 0.9
    
    return 0.1  # Default low confidence

@app.workflow("run_compliance_analysis", "Complete compliance gap analysis workflow")
async def run_compliance_analysis(workflow_state: WorkflowState):
    """Run comprehensive compliance gap analysis as a workflow"""
    try:
        # Step 1: Validation
        workflow_state.step = "validation"
        workflow_state.status = "processing"
        
        if not iso_summaries:
            workflow_state.status = "error"
            return {"error": "No ISO summary found. Upload ISO 13485:2024 Summary first."}
        
        if not clause_mappings:
            workflow_state.status = "error"
            return {"error": "No clause mappings found. Run clause mapping first."}
        
        await asyncio.sleep(1)
        
        # Step 2: Extract changed clauses
        workflow_state.step = "extracting_changes"
        
        latest_iso = list(iso_summaries.values())[-1]
        changed_clauses = set()
        
        for clause in latest_iso.new_clauses + latest_iso.modified_clauses:
            changed_clauses.add(clause.get('clause', ''))
        
        workflow_state.data["changed_clauses"] = list(changed_clauses)
        await asyncio.sleep(1)
        
        # Step 3: Gap analysis
        workflow_state.step = "gap_analysis"
        
        compliance_gaps.clear()
        mapped_clauses = {mapping.iso_clause for mapping in clause_mappings.values()}
        
        gaps_found = 0
        for changed_clause in changed_clauses:
            if not changed_clause:
                continue
            
            # Check if clause is covered by mappings
            found_mapping = any(changed_clause in clause for clause in mapped_clauses)
            
            if not found_mapping:
                gap = ComplianceGap(
                    id=str(uuid.uuid4()),
                    qsp_id="",
                    qsp_filename="Multiple",
                    iso_clause=changed_clause,
                    gap_type="missing",
                    description=f"No QSP documents address the new/modified clause: {changed_clause}",
                    severity="high",
                    recommendations=[
                        f"Create or update QSP documentation to address {changed_clause}",
                        "Review existing procedures for compliance gaps",
                        "Assign responsibility for implementing new requirements"
                    ]
                )
                compliance_gaps[gap.id] = gap
                gaps_found += 1
        
        workflow_state.data["gaps_found"] = gaps_found
        await asyncio.sleep(1)
        
        # Step 4: Score calculation
        workflow_state.step = "score_calculation"
        
        total_changed = len([c for c in changed_clauses if c])
        high_conf_mappings = sum(1 for m in clause_mappings.values() if m.confidence_score > 0.7)
        overall_score = min((high_conf_mappings / max(total_changed, 1)) * 100, 100)
        
        workflow_state.data["overall_score"] = overall_score
        
        # Step 5: Final results
        workflow_state.step = "compilation"
        
        result = {
            "workflow_id": workflow_state.workflow_id,
            "analysis_type": "ISO_13485_2024_compliance",
            "overall_score": round(overall_score, 2),
            "gaps_found": gaps_found,
            "high_priority_gaps": gaps_found,
            "total_documents": len(qsp_documents),
            "total_mappings": len(clause_mappings),
            "changed_clauses_analyzed": total_changed,
            "analysis_summary": {
                "compliant_areas": len(mapped_clauses),
                "gap_areas": gaps_found,
                "coverage_percentage": round((len(mapped_clauses) / max(total_changed, 1)) * 100, 1)
            },
            "completed_at": datetime.now(timezone.utc).isoformat()
        }
        
        # Store analysis in history
        analysis_history.append(result)
        
        workflow_state.step = "completed"
        workflow_state.status = "success"
        workflow_state.data["final_result"] = result
        
        return result
        
    except Exception as e:
        workflow_state.status = "error"
        return {"error": str(e), "workflow_id": workflow_state.workflow_id}

@app.session_action("save_analysis_session", "Save current analysis state to user session")
async def save_analysis_session(session: Session, session_name: str = "Default Analysis"):
    """Save current analysis state to user session"""
    
    analysis_state = {
        "session_name": session_name,
        "documents_count": len(qsp_documents),
        "iso_summaries_count": len(iso_summaries),
        "mappings_count": len(clause_mappings),
        "gaps_count": len(compliance_gaps),
        "saved_at": datetime.now(timezone.utc).isoformat(),
        "document_list": [{"id": doc.id, "filename": doc.filename} for doc in qsp_documents.values()],
        "latest_analysis": analysis_history[-1] if analysis_history else None
    }
    
    # Store in session
    if 'saved_analyses' not in session.data:
        session.data['saved_analyses'] = []
    
    session.data['saved_analyses'].append(analysis_state)
    session.set('last_saved_analysis', session_name)
    
    return {
        "status": "saved",
        "session_name": session_name,
        "analysis_state": analysis_state,
        "total_saved_sessions": len(session.data['saved_analyses']),
        "session_id": session.session_id
    }

@app.session_action("get_analysis_sessions", "Retrieve all saved analysis sessions")
async def get_analysis_sessions(session: Session):
    """Get all saved analysis sessions"""
    saved_analyses = session.get('saved_analyses', [])
    
    return {
        "saved_sessions": saved_analyses,
        "count": len(saved_analyses),
        "last_saved": session.get('last_saved_analysis'),
        "session_id": session.session_id
    }

@app.action("get_compliance_status", "Get current compliance overview with scores and key metrics")
def get_compliance_status(include_details: bool = False):
    """Get current compliance status and metrics"""
    
    latest_analysis = analysis_history[-1] if analysis_history else None
    
    status = {
        "system_status": "Ready" if qsp_documents and iso_summaries else "Setup Required",
        "total_documents": len(qsp_documents),
        "total_mappings": len(clause_mappings),
        "total_gaps": len(compliance_gaps),
        "iso_summaries_loaded": len(iso_summaries),
        "overall_score": latest_analysis["overall_score"] if latest_analysis else None,
        "last_analysis_date": latest_analysis["completed_at"] if latest_analysis else None,
        "analysis_ready": len(qsp_documents) > 0 and len(iso_summaries) > 0,
        "status_timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    if include_details and latest_analysis:
        status["latest_analysis_details"] = latest_analysis
        status["gap_breakdown"] = {
            "high": len([g for g in compliance_gaps.values() if g.severity == "high"]),
            "medium": len([g for g in compliance_gaps.values() if g.severity == "medium"]),
            "low": len([g for g in compliance_gaps.values() if g.severity == "low"])
        }
    
    return status

@app.action("get_detailed_gaps", "Get detailed compliance gaps with recommendations")
def get_detailed_gaps(severity: str = "all", limit: int = 10):
    """Get detailed compliance gaps with recommendations"""
    
    gaps = list(compliance_gaps.values())
    
    if severity != "all":
        gaps = [gap for gap in gaps if gap.severity == severity]
    
    # Limit results
    gaps = gaps[:limit]
    
    if not gaps:
        return {
            "gaps": [],
            "count": 0,
            "message": "No compliance gaps found. Excellent compliance status!",
            "severity_filter": severity
        }
    
    return {
        "gaps": [gap.dict() for gap in gaps],
        "count": len(gaps),
        "total_gaps": len(compliance_gaps),
        "severity_filter": severity,
        "message": f"Found {len(gaps)} compliance gaps (showing top {limit})"
    }

@app.action("query_specific_clause", "Query compliance status for a specific ISO clause")
def query_specific_clause(clause: str, include_recommendations: bool = True):
    """Query compliance for a specific ISO clause"""
    
    # Find mappings for this clause
    related_mappings = [
        mapping for mapping in clause_mappings.values()
        if clause.lower() in mapping.iso_clause.lower()
    ]
    
    # Find gaps for this clause
    related_gaps = [
        gap for gap in compliance_gaps.values()
        if clause.lower() in gap.iso_clause.lower()
    ]
    
    result = {
        "clause_queried": clause,
        "mappings_found": len(related_mappings),
        "gaps_found": len(related_gaps),
        "compliance_status": "compliant" if related_mappings and not related_gaps else "non-compliant",
        "analysis_timestamp": datetime.now(timezone.utc).isoformat()
    }
    
    if related_mappings:
        result["mapped_documents"] = [
            {
                "filename": mapping.qsp_filename,
                "section": mapping.section_title,
                "confidence": mapping.confidence_score,
                "evidence_preview": mapping.evidence_text[:100] + "..."
            }
            for mapping in related_mappings[:3]  # Top 3
        ]
    
    if related_gaps and include_recommendations:
        result["compliance_gaps"] = [
            {
                "gap_id": gap.id,
                "gap_type": gap.gap_type,
                "description": gap.description,
                "severity": gap.severity,
                "recommendations": gap.recommendations[:3]  # Top 3 recommendations
            }
            for gap in related_gaps[:2]  # Top 2 gaps
        ]
    
    return result

@app.action("get_analysis_history", "Get history of all compliance analyses performed")
def get_analysis_history(limit: int = 5):
    """Get history of compliance analyses"""
    
    recent_analyses = analysis_history[-limit:] if limit > 0 else analysis_history
    recent_analyses.reverse()  # Most recent first
    
    return {
        "analyses": recent_analyses,
        "count": len(recent_analyses),
        "total_analyses_performed": len(analysis_history),
        "message": f"Showing {len(recent_analyses)} most recent analyses"
    }

# Enable authentication (optional - uncomment to enable)
# app.enable_auth("your-secret-jwt-key-here")

if __name__ == "__main__":
    print("🚀 QSP Compliance Checker - Kalibr SDK v2")
    print("=" * 50)
    print("Multi-Model AI Integration Ready!")
    print("")
    print("📊 Features Available:")
    print("✅ Upload QSP documents (.docx, .txt)")
    print("✅ Upload ISO 13485:2024 summaries")
    print("✅ AI-powered clause mapping with streaming")
    print("✅ Comprehensive compliance analysis workflow")
    print("✅ Session-based analysis state management")
    print("✅ Detailed gap analysis with recommendations")
    print("")
    print("🤖 AI Models Supported:")
    print("✅ ChatGPT (GPT Actions)")
    print("✅ Claude (MCP Protocol)")
    print("✅ Google Gemini (Extensions)")
    print("✅ Microsoft Copilot (Plugins)")
    print("")
    print("🔧 Run with:")
    print("   kalibr-connect serve qsp_compliance_kalibr_v2.py")
    print("")
    print("📡 Endpoints auto-generated:")
    print("   http://localhost:8000/gpt-actions.json    # For ChatGPT")
    print("   http://localhost:8000/mcp.json            # For Claude")
    print("   http://localhost:8000/schemas/gemini      # For Gemini")
    print("   http://localhost:8000/schemas/copilot     # For Copilot")
    print("   http://localhost:8000/docs                # Interactive API docs")