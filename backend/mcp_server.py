#!/usr/bin/env python3
"""
QSP Compliance Checker MCP Server

This MCP server provides Claude/ChatGPT with tools to analyze QSP documents
for ISO 13485:2024 compliance. It runs alongside the web UI, providing
dual interface access.
"""

import asyncio
import json
import os
import sys
import base64
import io
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from datetime import datetime, timezone
import uuid

import aiofiles
from docx import Document
from dotenv import load_dotenv
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field

from mcp.server import Server
from mcp.server.models import InitializationOptions
from mcp.server.stdio import stdio_server
from mcp.types import (
    Resource,
    Tool,
    TextContent,
    ImageContent,
    EmbeddedResource,
    LoggingLevel
)

from emergentintegrations.llm.chat import LlmChat, UserMessage

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("qsp-compliance-mcp")

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'test_database')]

# MCP Server instance
server = Server("qsp-compliance-checker")

# Data Models (reusing from main server)
class QSPDocument(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    content: str
    sections: Dict[str, str] = Field(default_factory=dict)
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    processed: bool = False

class ISOSummary(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    framework: str = "ISO_13485"
    version: str = "2024_summary"
    content: str
    new_clauses: List[Dict[str, str]] = Field(default_factory=list)
    modified_clauses: List[Dict[str, str]] = Field(default_factory=list)
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ClauseMapping(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    qsp_id: str
    qsp_filename: str
    section_title: str
    section_content: str
    iso_clause: str
    confidence_score: float
    evidence_text: str
    mapping_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ComplianceGap(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    qsp_id: str
    qsp_filename: str
    iso_clause: str
    gap_type: str
    description: str
    severity: str
    recommendations: List[str] = Field(default_factory=list)

class ComplianceAnalysis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    overall_score: float
    total_documents: int
    compliant_documents: int
    gaps: List[ComplianceGap]
    affected_documents: List[str]
    analysis_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Utility Functions
def prepare_for_mongo(data):
    """Convert datetime objects to ISO strings for MongoDB storage"""
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, datetime):
                data[key] = value.isoformat()
            elif isinstance(value, dict):
                data[key] = prepare_for_mongo(value)
            elif isinstance(value, list):
                data[key] = [prepare_for_mongo(item) if isinstance(item, dict) else item for item in value]
    return data

def parse_from_mongo(item):
    """Convert ISO string timestamps back to datetime objects"""
    if isinstance(item, dict):
        for key, value in item.items():
            if key.endswith('_date') and isinstance(value, str):
                try:
                    item[key] = datetime.fromisoformat(value)
                except:
                    pass
            elif isinstance(value, dict):
                item[key] = parse_from_mongo(value)
            elif isinstance(value, list):
                item[key] = [parse_from_mongo(sub_item) if isinstance(sub_item, dict) else sub_item for sub_item in value]
    return item

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract text from DOCX file: {e}")

async def parse_document_sections(content: str, filename: str) -> Dict[str, str]:
    """Parse document content into logical sections"""
    sections = {}
    
    # Split by common section patterns
    import re
    patterns = [
        r'\b\d+\.\s+[A-Z][^.]*\n',  # Numbered sections
        r'\b[A-Z][^.]{10,50}:?\n',   # Header patterns
        r'^[A-Z\s]{3,}\n',           # All caps headers
    ]
    
    current_section = "Introduction"
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line is a section header
        is_header = False
        for pattern in patterns:
            if re.match(pattern, line + '\n'):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = line.replace(':', '').strip()
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(line)
    
    # Save final section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    # If no sections found, treat entire document as one section
    if not sections:
        sections[f"{filename} - Full Content"] = content
    
    return sections

async def analyze_clause_mapping(qsp_content: str, qsp_filename: str, iso_clauses: List[str]) -> List[Dict[str, Any]]:
    """Use AI to map QSP content to ISO clauses"""
    try:
        # Initialize LLM chat
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=f"mapping_{uuid.uuid4()}",
            system_message="""You are an expert in ISO 13485 compliance for medical devices. 
            Analyze QSP document sections and map them to relevant ISO 13485:2024 clauses.
            
            For each mapping, provide:
            1. The ISO clause number and title
            2. Confidence score (0.0-1.0) based on relevance
            3. Evidence text from the QSP that supports the mapping
            4. Brief explanation of the mapping
            
            Return results as a JSON array with this structure:
            [
                {
                    "iso_clause": "4.1 General requirements",
                    "confidence_score": 0.85,
                    "evidence_text": "relevant excerpt from QSP",
                    "explanation": "brief explanation of why this maps"
                }
            ]"""
        ).with_model("openai", "gpt-4o")
        
        iso_clauses_text = "\n".join([f"- {clause}" for clause in iso_clauses])
        
        user_message = UserMessage(
            text=f"""Analyze this QSP document content and map it to relevant ISO 13485:2024 clauses:

QSP Document: {qsp_filename}
Content:
{qsp_content[:3000]}...

ISO 13485:2024 Key Clauses to consider:
{iso_clauses_text}

Provide detailed mappings with confidence scores."""
        )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            # Clean the response - remove markdown formatting if present
            clean_response = response.strip()
            if clean_response.startswith('```json'):
                clean_response = clean_response[7:]  # Remove ```json
            if clean_response.endswith('```'):
                clean_response = clean_response[:-3]  # Remove ```
            clean_response = clean_response.strip()
            
            mappings = json.loads(clean_response)
            return mappings if isinstance(mappings, list) else []
        except json.JSONDecodeError:
            logger.error(f"Failed to parse AI response as JSON: {response}")
            return []
            
    except Exception as e:
        logger.error(f"Error in AI clause mapping: {e}")
        return []

def format_response_adaptively(data: Any, request_type: str, detail_level: str = "auto") -> str:
    """Format responses adaptively based on request context"""
    
    if detail_level == "json" or request_type.lower().endswith("_json"):
        return json.dumps(data, indent=2, default=str)
    
    if isinstance(data, dict):
        if "compliance_score" in data:  # Dashboard data
            return f"""📊 **QSP Compliance Status**
            
**Overall Compliance Score:** {data['compliance_score']}%
**Documents Analyzed:** {data['total_documents']}
**Compliance Gaps:** {data['gaps_count']} issues identified
**AI Clause Mappings:** {data['total_mappings']} generated
**ISO Summary:** {'✅ Loaded' if data['iso_summary_loaded'] else '❌ Not uploaded'}

Last Analysis: {data['last_analysis_date'] or 'Never'}"""
        
        elif "gaps" in data:  # Gaps data
            if not data["gaps"]:
                return "✅ **No Compliance Gaps Found**\n\nExcellent! Your QSP documents appear to be fully compliant with ISO 13485:2024 changes."
            
            gaps_text = f"⚠️ **{len(data['gaps'])} Compliance Gaps Identified**\n\n"
            for i, gap in enumerate(data["gaps"][:5], 1):  # Show first 5 gaps
                gaps_text += f"**{i}. {gap['iso_clause']}** ({gap['severity'].upper()} priority)\n"
                gaps_text += f"   📋 Document: {gap['qsp_filename']}\n"
                gaps_text += f"   📝 Issue: {gap['description']}\n\n"
            
            if len(data["gaps"]) > 5:
                gaps_text += f"... and {len(data['gaps']) - 5} more gaps\n"
            
            return gaps_text
    
    elif isinstance(data, list):
        if len(data) > 0 and isinstance(data[0], dict):
            if "filename" in data[0]:  # Documents list
                docs_text = f"📄 **{len(data)} QSP Documents**\n\n"
                for doc in data:
                    docs_text += f"• **{doc['filename']}** - {len(doc.get('sections', {}))} sections\n"
                    docs_text += f"  Uploaded: {doc['upload_date']}\n\n"
                return docs_text
            
            elif "iso_clause" in data[0] and "confidence_score" in data[0]:  # Mappings
                mappings_text = f"🔗 **{len(data)} Clause Mappings**\n\n"
                for mapping in data[:10]:  # Show first 10
                    mappings_text += f"• **{mapping['iso_clause']}** (confidence: {mapping['confidence_score']:.2f})\n"
                    mappings_text += f"  📄 {mapping['qsp_filename']} - {mapping['section_title']}\n"
                    mappings_text += f"  📝 Evidence: {mapping['evidence_text'][:100]}...\n\n"
                
                if len(data) > 10:
                    mappings_text += f"... and {len(data) - 10} more mappings\n"
                
                return mappings_text
    
    # Fallback to JSON for complex data
    return json.dumps(data, indent=2, default=str)

# MCP Tool Definitions
@server.list_tools()
async def handle_list_tools() -> list[Tool]:
    """Return the list of available tools"""
    return [
        # Document Management Tools
        Tool(
            name="upload_qsp_document",
            description="Upload a QSP (Quality System Procedure) document for compliance analysis. Accepts file path or base64 content.",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Name of the file"},
                    "file_path": {"type": "string", "description": "Path to the file (optional if base64_content provided)"},
                    "base64_content": {"type": "string", "description": "Base64 encoded file content (optional if file_path provided)"},
                    "file_type": {"type": "string", "enum": ["txt", "docx"], "description": "File type: txt or docx"}
                },
                "required": ["filename", "file_type"],
                "oneOf": [
                    {"required": ["file_path"]},
                    {"required": ["base64_content"]}
                ]
            }
        ),
        
        Tool(
            name="upload_iso_summary",
            description="Upload ISO 13485:2024 Summary of Changes document for compliance baseline.",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Name of the file"},
                    "file_path": {"type": "string", "description": "Path to the file (optional if base64_content provided)"},
                    "base64_content": {"type": "string", "description": "Base64 encoded file content (optional if file_path provided)"},
                    "file_type": {"type": "string", "enum": ["txt", "docx", "pdf"], "description": "File type"}
                },
                "required": ["filename", "file_type"],
                "oneOf": [
                    {"required": ["file_path"]},
                    {"required": ["base64_content"]}
                ]
            }
        ),
        
        Tool(
            name="list_documents",
            description="List all uploaded QSP documents with their processing status.",
            inputSchema={
                "type": "object",
                "properties": {
                    "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                }
            }
        ),
        
        # Analysis Tools
        Tool(
            name="run_clause_mapping",
            description="Run AI-powered analysis to map QSP document sections to ISO 13485 clauses.",
            inputSchema={
                "type": "object",
                "properties": {
                    "document_ids": {"type": "array", "items": {"type": "string"}, "description": "Specific document IDs to analyze (optional - analyzes all if not provided)"}
                }
            }
        ),
        
        Tool(
            name="run_compliance_analysis",
            description="Run comprehensive compliance gap analysis against ISO 13485:2024 changes.",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        
        Tool(
            name="get_compliance_status",
            description="Get current overall compliance status and key metrics.",
            inputSchema={
                "type": "object",
                "properties": {
                    "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                }
            }
        ),
        
        # Reporting Tools
        Tool(
            name="get_dashboard_summary",
            description="Get dashboard overview with compliance scores, document counts, and analysis status.",
            inputSchema={
                "type": "object",
                "properties": {
                    "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                }
            }
        ),
        
        Tool(
            name="get_detailed_gaps",
            description="Get detailed compliance gaps with recommendations and priority levels.",
            inputSchema={
                "type": "object",
                "properties": {
                    "severity": {"type": "string", "enum": ["all", "high", "medium", "low"], "default": "all"},
                    "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                }
            }
        ),
        
        Tool(
            name="get_clause_mappings",
            description="Get AI-generated clause mappings between QSP documents and ISO clauses.",
            inputSchema={
                "type": "object",
                "properties": {
                    "document_id": {"type": "string", "description": "Filter by specific document ID (optional)"},
                    "iso_clause": {"type": "string", "description": "Filter by specific ISO clause (optional)"},
                    "min_confidence": {"type": "number", "minimum": 0, "maximum": 1, "description": "Minimum confidence score (optional)"},
                    "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                }
            }
        ),
        
        Tool(
            name="query_specific_clause",
            description="Query compliance status for a specific ISO 13485 clause across all QSP documents.",
            inputSchema={
                "type": "object",
                "properties": {
                    "clause": {"type": "string", "description": "ISO clause to query (e.g., '4.1.6', '7.3.10')"},
                    "include_recommendations": {"type": "boolean", "default": True}
                }
            }
        )
    ]

# Tool Handlers
@server.call_tool()
async def handle_call_tool(name: str, arguments: dict) -> list[TextContent]:
    """Handle tool calls"""
    
    try:
        if name == "upload_qsp_document":
            return await handle_upload_qsp_document(arguments)
        elif name == "upload_iso_summary":
            return await handle_upload_iso_summary(arguments)
        elif name == "list_documents":
            return await handle_list_documents(arguments)
        elif name == "run_clause_mapping":
            return await handle_run_clause_mapping(arguments)
        elif name == "run_compliance_analysis":
            return await handle_run_compliance_analysis(arguments)
        elif name == "get_compliance_status":
            return await handle_get_compliance_status(arguments)
        elif name == "get_dashboard_summary":
            return await handle_get_dashboard_summary(arguments)
        elif name == "get_detailed_gaps":
            return await handle_get_detailed_gaps(arguments)
        elif name == "get_clause_mappings":
            return await handle_get_clause_mappings(arguments)
        elif name == "query_specific_clause":
            return await handle_query_specific_clause(arguments)
        else:
            return [TextContent(type="text", text=f"Unknown tool: {name}")]
            
    except Exception as e:
        logger.error(f"Error in tool {name}: {e}")
        return [TextContent(type="text", text=f"Error executing {name}: {str(e)}")]

# Tool Implementation Functions
async def handle_upload_qsp_document(args: dict) -> list[TextContent]:
    """Handle QSP document upload"""
    filename = args["filename"]
    file_type = args["file_type"]
    
    # Get file content
    if "file_path" in args:
        file_path = Path(args["file_path"])
        if not file_path.exists():
            return [TextContent(type="text", text=f"File not found: {file_path}")]
        
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
    else:
        # Decode base64 content
        try:
            content = base64.b64decode(args["base64_content"])
        except Exception as e:
            return [TextContent(type="text", text=f"Failed to decode base64 content: {e}")]
    
    # Extract text
    try:
        if file_type == "docx":
            text_content = await extract_text_from_docx(content)
        else:  # txt
            text_content = content.decode('utf-8')
    except Exception as e:
        return [TextContent(type="text", text=f"Failed to extract text: {e}")]
    
    # Parse into sections
    sections = await parse_document_sections(text_content, filename)
    
    # Create QSP document
    qsp_doc = QSPDocument(
        filename=filename,
        content=text_content,
        sections=sections,
        processed=True
    )
    
    # Store in MongoDB
    doc_dict = prepare_for_mongo(qsp_doc.model_dump())
    await db.qsp_documents.insert_one(doc_dict)
    
    logger.info(f"Uploaded QSP document via MCP: {filename}")
    
    response = f"""✅ **QSP Document Uploaded Successfully**

**File:** {filename}
**Document ID:** {qsp_doc.id}
**Sections Parsed:** {len(sections)}
**Content Length:** {len(text_content)} characters

The document has been processed and is ready for clause mapping analysis."""

    return [TextContent(type="text", text=response)]

async def handle_upload_iso_summary(args: dict) -> list[TextContent]:
    """Handle ISO summary upload"""
    filename = args["filename"]
    file_type = args["file_type"]
    
    # Get file content
    if "file_path" in args:
        file_path = Path(args["file_path"])
        if not file_path.exists():
            return [TextContent(type="text", text=f"File not found: {file_path}")]
        
        async with aiofiles.open(file_path, 'rb') as f:
            content = await f.read()
    else:
        # Decode base64 content
        try:
            content = base64.b64decode(args["base64_content"])
        except Exception as e:
            return [TextContent(type="text", text=f"Failed to decode base64 content: {e}")]
    
    # Extract text
    try:
        if file_type == "docx":
            text_content = await extract_text_from_docx(content)
        else:  # txt or pdf
            text_content = content.decode('utf-8')
    except Exception as e:
        return [TextContent(type="text", text=f"Failed to extract text: {e}")]
    
    # Parse ISO summary content
    new_clauses = []
    modified_clauses = []
    
    # Simple parsing for new and modified clauses
    lines = text_content.split('\n')
    current_section = ""
    
    import re
    for line in lines:
        line = line.strip()
        if 'NEW CLAUSES' in line.upper() or 'NEW REQUIREMENTS' in line.upper():
            current_section = "new"
        elif 'MODIFIED CLAUSES' in line.upper() or 'CHANGED CLAUSES' in line.upper():
            current_section = "modified"
        elif line and current_section and re.match(r'\d+\.\d+', line):
            clause_info = {"clause": line, "description": line}
            if current_section == "new":
                new_clauses.append(clause_info)
            elif current_section == "modified":
                modified_clauses.append(clause_info)
    
    # Create ISO summary
    iso_summary = ISOSummary(
        content=text_content,
        new_clauses=new_clauses,
        modified_clauses=modified_clauses
    )
    
    # Store in MongoDB
    summary_dict = prepare_for_mongo(iso_summary.model_dump())
    await db.iso_summaries.insert_one(summary_dict)
    
    logger.info(f"Uploaded ISO summary via MCP: {filename}")
    
    response = f"""✅ **ISO 13485:2024 Summary Uploaded Successfully**

**File:** {filename}
**Summary ID:** {iso_summary.id}
**New Clauses Identified:** {len(new_clauses)}
**Modified Clauses Identified:** {len(modified_clauses)}

The ISO summary has been processed and compliance analysis is now enabled."""

    return [TextContent(type="text", text=response)]

async def handle_list_documents(args: dict) -> list[TextContent]:
    """Handle list documents request"""
    format_type = args.get("format", "summary")
    
    # Get all QSP documents
    docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
    docs = [parse_from_mongo(doc) for doc in docs]
    
    response = format_response_adaptively(docs, "list_documents", format_type)
    return [TextContent(type="text", text=response)]

async def handle_run_clause_mapping(args: dict) -> list[TextContent]:
    """Handle run clause mapping request"""
    document_ids = args.get("document_ids", [])
    
    # Get QSP documents
    query = {"id": {"$in": document_ids}} if document_ids else {}
    qsp_docs = await db.qsp_documents.find(query, {"_id": 0}).to_list(length=None)
    
    if not qsp_docs:
        return [TextContent(type="text", text="❌ No QSP documents found. Upload documents first.")]
    
    # Get ISO summary for clauses
    iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
    if not iso_summary:
        return [TextContent(type="text", text="❌ No ISO summary found. Upload ISO 13485:2024 Summary first.")]
    
    # Standard ISO 13485 clauses
    iso_clauses = [
        "4.1 General requirements", "4.2 Documentation requirements",
        "5.1 Management commitment", "5.2 Customer focus", "5.3 Quality policy",
        "5.4 Planning", "5.5 Responsibility, authority and communication",
        "5.6 Management review", "6.1 Provision of resources", "6.2 Human resources",
        "6.3 Infrastructure", "6.4 Work environment", "7.1 Planning of product realization",
        "7.2 Customer-related processes", "7.3 Design and development", "7.4 Purchasing",
        "7.5 Production and service provision", "7.6 Control of monitoring and measuring equipment",
        "8.1 General", "8.2 Monitoring and measurement", "8.3 Control of nonconforming product",
        "8.4 Analysis of data", "8.5 Improvement"
    ]
    
    # Clear existing mappings for these documents
    if document_ids:
        await db.clause_mappings.delete_many({"qsp_id": {"$in": document_ids}})
    else:
        await db.clause_mappings.delete_many({})
    
    total_mappings = 0
    
    # Process each QSP document
    for qsp_doc in qsp_docs:
        qsp_doc = parse_from_mongo(qsp_doc)
        
        # Map each section
        for section_title, section_content in qsp_doc['sections'].items():
            if len(section_content) < 50:  # Skip very short sections
                continue
            
            mappings = await analyze_clause_mapping(
                section_content, qsp_doc['filename'], iso_clauses
            )
            
            for mapping in mappings:
                if mapping.get('confidence_score', 0) > 0.3:  # Only store high-confidence mappings
                    clause_mapping = ClauseMapping(
                        qsp_id=qsp_doc['id'],
                        qsp_filename=qsp_doc['filename'],
                        section_title=section_title,
                        section_content=section_content[:500],  # Truncate for storage
                        iso_clause=mapping['iso_clause'],
                        confidence_score=mapping['confidence_score'],
                        evidence_text=mapping['evidence_text']
                    )
                    
                    mapping_dict = prepare_for_mongo(clause_mapping.model_dump())
                    await db.clause_mappings.insert_one(mapping_dict)
                    total_mappings += 1
    
    logger.info(f"Generated {total_mappings} clause mappings via MCP")
    
    response = f"""🤖 **AI Clause Mapping Completed**

**Documents Processed:** {len(qsp_docs)}
**Mappings Generated:** {total_mappings}
**AI Model:** GPT-4o with ISO 13485 expertise

The AI has analyzed your QSP documents and mapped sections to relevant ISO clauses with confidence scores. You can now run compliance analysis."""

    return [TextContent(type="text", text=response)]

async def handle_run_compliance_analysis(args: dict) -> list[TextContent]:
    """Handle run compliance analysis request"""
    # Get ISO summary and mappings
    iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
    if not iso_summary:
        return [TextContent(type="text", text="❌ No ISO summary found. Upload ISO 13485:2024 Summary first.")]
    
    mappings = await db.clause_mappings.find({}, {"_id": 0}).to_list(length=None)
    qsp_docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
    
    if not mappings:
        return [TextContent(type="text", text="❌ No clause mappings found. Run clause mapping first.")]
    
    # Extract changed clauses from ISO summary
    changed_clauses = set()
    iso_summary = parse_from_mongo(iso_summary)
    
    for clause in iso_summary.get('new_clauses', []):
        changed_clauses.add(clause.get('clause', ''))
    for clause in iso_summary.get('modified_clauses', []):
        changed_clauses.add(clause.get('clause', ''))
    
    # Find gaps
    gaps = []
    mapped_clauses = {}
    
    # Group mappings by clause
    for mapping in mappings:
        clause = mapping['iso_clause']
        if clause not in mapped_clauses:
            mapped_clauses[clause] = []
        mapped_clauses[clause].append(mapping)
    
    # Check for gaps in changed clauses
    for changed_clause in changed_clauses:
        if not changed_clause:
            continue
            
        # Find if any existing mappings cover this clause
        found_mapping = False
        low_confidence = []
        
        for clause, clause_mappings in mapped_clauses.items():
            if changed_clause in clause or any(changed_clause in m['iso_clause'] for m in clause_mappings):
                found_mapping = True
                # Check confidence scores
                for mapping in clause_mappings:
                    if mapping['confidence_score'] < 0.7:
                        low_confidence.append(mapping)
        
        if not found_mapping:
            gap = ComplianceGap(
                qsp_id="",
                qsp_filename="Multiple",
                iso_clause=changed_clause,
                gap_type="missing",
                description=f"No QSP documents found addressing the new/modified clause: {changed_clause}",
                severity="high",
                recommendations=[
                    f"Create or update QSP documentation to address {changed_clause}",
                    "Review existing procedures for compliance gaps",
                    "Assign responsibility for implementing new requirements"
                ]
            )
            gaps.append(gap)
        elif low_confidence:
            for mapping in low_confidence:
                gap = ComplianceGap(
                    qsp_id=mapping['qsp_id'],
                    qsp_filename=mapping['qsp_filename'],
                    iso_clause=mapping['iso_clause'],
                    gap_type="partial",
                    description=f"Low confidence mapping for {mapping['iso_clause']} in {mapping['qsp_filename']}",
                    severity="medium",
                    recommendations=[
                        "Review and strengthen documentation",
                        "Add more specific requirements",
                        "Clarify procedures and responsibilities"
                    ]
                )
                gaps.append(gap)
    
    # Calculate compliance score
    total_changed_clauses = len([c for c in changed_clauses if c])
    high_confidence_mappings = sum(1 for mapping in mappings if mapping['confidence_score'] > 0.7)
    
    if total_changed_clauses > 0:
        overall_score = min(high_confidence_mappings / total_changed_clauses * 100, 100)
    else:
        overall_score = 100
    
    # Get affected documents
    affected_docs = list(set(gap.qsp_filename for gap in gaps if gap.qsp_filename != "Multiple"))
    
    # Create analysis
    analysis = ComplianceAnalysis(
        overall_score=round(overall_score, 2),
        total_documents=len(qsp_docs),
        compliant_documents=len(qsp_docs) - len(affected_docs),
        gaps=gaps,
        affected_documents=affected_docs
    )
    
    # Store analysis
    analysis_dict = prepare_for_mongo(analysis.model_dump())
    await db.compliance_analyses.delete_many({})  # Keep only latest
    await db.compliance_analyses.insert_one(analysis_dict)
    
    logger.info(f"Compliance analysis completed via MCP. Found {len(gaps)} gaps.")
    
    high_gaps = len([g for g in gaps if g.severity == "high"])
    medium_gaps = len([g for g in gaps if g.severity == "medium"])
    low_gaps = len([g for g in gaps if g.severity == "low"])
    
    response = f"""📊 **Compliance Analysis Completed**

**Overall Compliance Score:** {analysis.overall_score}%
**Documents Analyzed:** {analysis.total_documents}
**Compliant Documents:** {analysis.compliant_documents}

**Gaps Identified:**
• 🔴 High Priority: {high_gaps}
• 🟡 Medium Priority: {medium_gaps} 
• 🔵 Low Priority: {low_gaps}

**Affected Documents:** {len(affected_docs)} documents need updates

The analysis compared your QSP documents against ISO 13485:2024 changes and identified specific areas requiring attention."""

    return [TextContent(type="text", text=response)]

async def handle_get_compliance_status(args: dict) -> list[TextContent]:
    """Handle get compliance status request"""
    format_type = args.get("format", "summary")
    
    # Get latest analysis
    analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
    
    # Get document counts
    total_docs = await db.qsp_documents.count_documents({})
    total_mappings = await db.clause_mappings.count_documents({})
    
    # Get ISO summary
    iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
    
    if analysis:
        analysis = parse_from_mongo(analysis)
    
    if iso_summary:
        iso_summary = parse_from_mongo(iso_summary)
    
    status_data = {
        "compliance_score": analysis['overall_score'] if analysis else 0,
        "total_documents": total_docs,
        "total_mappings": total_mappings,
        "gaps_count": len(analysis['gaps']) if analysis else 0,
        "affected_documents": len(analysis['affected_documents']) if analysis else 0,
        "iso_summary_loaded": iso_summary is not None,
        "new_clauses_count": len(iso_summary.get('new_clauses', [])) if iso_summary else 0,
        "modified_clauses_count": len(iso_summary.get('modified_clauses', [])) if iso_summary else 0,
        "last_analysis_date": analysis['analysis_date'] if analysis else None
    }
    
    response = format_response_adaptively(status_data, "compliance_status", format_type)
    return [TextContent(type="text", text=response)]

async def handle_get_dashboard_summary(args: dict) -> list[TextContent]:
    """Handle get dashboard summary request"""
    # Reuse compliance status logic
    return await handle_get_compliance_status(args)

async def handle_get_detailed_gaps(args: dict) -> list[TextContent]:
    """Handle get detailed gaps request"""
    severity_filter = args.get("severity", "all")
    format_type = args.get("format", "summary")
    
    analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
    
    if not analysis:
        return [TextContent(type="text", text="❌ No compliance analysis found. Run analysis first.")]
    
    analysis = parse_from_mongo(analysis)
    gaps = analysis['gaps']
    
    # Filter by severity
    if severity_filter != "all":
        gaps = [gap for gap in gaps if gap['severity'] == severity_filter]
    
    gaps_data = {
        "gaps": gaps,
        "total_gaps": len(gaps),
        "analysis_date": analysis['analysis_date']
    }
    
    response = format_response_adaptively(gaps_data, "detailed_gaps", format_type)
    return [TextContent(type="text", text=response)]

async def handle_get_clause_mappings(args: dict) -> list[TextContent]:
    """Handle get clause mappings request"""
    document_id = args.get("document_id")
    iso_clause = args.get("iso_clause") 
    min_confidence = args.get("min_confidence", 0.0)
    format_type = args.get("format", "summary")
    
    # Build query
    query = {}
    if document_id:
        query["qsp_id"] = document_id
    if iso_clause:
        query["iso_clause"] = {"$regex": iso_clause, "$options": "i"}
    if min_confidence > 0:
        query["confidence_score"] = {"$gte": min_confidence}
    
    mappings = await db.clause_mappings.find(query, {"_id": 0}).to_list(length=None)
    mappings = [parse_from_mongo(mapping) for mapping in mappings]
    
    response = format_response_adaptively(mappings, "clause_mappings", format_type)
    return [TextContent(type="text", text=response)]

async def handle_query_specific_clause(args: dict) -> list[TextContent]:
    """Handle query specific clause request"""
    clause = args["clause"]
    include_recommendations = args.get("include_recommendations", True)
    
    # Find mappings for this clause
    mappings = await db.clause_mappings.find(
        {"iso_clause": {"$regex": clause, "$options": "i"}},
        {"_id": 0}
    ).to_list(length=None)
    
    # Find gaps for this clause
    analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
    relevant_gaps = []
    
    if analysis:
        analysis = parse_from_mongo(analysis)
        relevant_gaps = [gap for gap in analysis['gaps'] if clause in gap['iso_clause']]
    
    if not mappings and not relevant_gaps:
        response = f"❓ **No Coverage Found for Clause {clause}**\n\nThis clause is not currently addressed in your QSP documents. Consider creating new documentation to cover these requirements."
    else:
        response = f"🔍 **Analysis for ISO Clause {clause}**\n\n"
        
        if mappings:
            response += f"**📋 QSP Coverage ({len(mappings)} mappings found):**\n"
            for mapping in mappings[:5]:  # Show first 5
                response += f"• **{mapping['qsp_filename']}** - {mapping['section_title']}\n"
                response += f"  Confidence: {mapping['confidence_score']:.2f}\n"
                response += f"  Evidence: {mapping['evidence_text'][:100]}...\n\n"
        
        if relevant_gaps:
            response += f"**⚠️ Compliance Gaps ({len(relevant_gaps)} found):**\n"
            for gap in relevant_gaps:
                response += f"• **{gap['severity'].upper()} Priority:** {gap['description']}\n"
                if include_recommendations and gap['recommendations']:
                    response += f"  Recommendations:\n"
                    for rec in gap['recommendations'][:3]:
                        response += f"    - {rec}\n"
                response += "\n"
    
    return [TextContent(type="text", text=response)]

# Main server function
async def main():
    """Main server function"""
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="qsp-compliance-checker",
                server_version="1.0.0",
                capabilities=server.get_capabilities(
                    notification_options=None,
                    experimental_capabilities={}
                )
            )
        )

if __name__ == "__main__":
    asyncio.run(main())