from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form, WebSocket, Depends
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import json
import io
import re
import base64
from docx import Document
from emergentintegrations.llm.chat import LlmChat, UserMessage

# Import multi-tenant components
from api import auth as auth_router_module
from core.auth import get_current_user, get_current_user_optional
from core.storage_service import storage_service

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI(title="QSP Compliance Checker", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Pydantic Models - Multi-tenant aware
class QSPDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tenant_id: str  # Multi-tenant isolation
    filename: str
    content: str
    sections: Dict[str, str] = Field(default_factory=dict)
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    processed: bool = False
    uploaded_by: Optional[str] = None  # user_id

class ISOSummary(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tenant_id: str  # Multi-tenant isolation
    framework: str = "ISO_13485"
    version: str = "2024_summary"
    content: str
    new_clauses: List[Dict[str, str]] = Field(default_factory=list)
    modified_clauses: List[Dict[str, str]] = Field(default_factory=list)
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    uploaded_by: Optional[str] = None  # user_id

class ClauseMapping(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tenant_id: str  # Multi-tenant isolation
    qsp_id: str
    qsp_filename: str
    section_title: str
    section_content: str
    iso_clause: str
    confidence_score: float
    evidence_text: str
    mapping_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ComplianceGap(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tenant_id: str  # Multi-tenant isolation
    qsp_id: str
    qsp_filename: str
    iso_clause: str
    gap_type: str  # "missing", "partial", "outdated"
    description: str
    severity: str  # "high", "medium", "low"
    recommendations: List[str] = Field(default_factory=list)

class ComplianceAnalysis(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tenant_id: str  # Multi-tenant isolation
    overall_score: float
    total_documents: int
    compliant_documents: int
    gaps: List[ComplianceGap]
    affected_documents: List[str]
    analysis_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Helper Functions
def prepare_for_mongo(data):
    """Convert datetime objects to ISO strings for MongoDB storage"""
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, datetime):
                data[key] = value.isoformat()
            elif isinstance(value, dict):
                data[key] = prepare_for_mongo(value)
            elif isinstance(value, list):
                data[key] = [prepare_for_mongo(item) if isinstance(item, dict) else item for item in value]
    return data

def parse_from_mongo(item):
    """Convert ISO string timestamps back to datetime objects"""
    if isinstance(item, dict):
        for key, value in item.items():
            if key.endswith('_date') and isinstance(value, str):
                try:
                    item[key] = datetime.fromisoformat(value)
                except:
                    pass
            elif isinstance(value, dict):
                item[key] = parse_from_mongo(value)
            elif isinstance(value, list):
                item[key] = [parse_from_mongo(sub_item) if isinstance(sub_item, dict) else sub_item for sub_item in value]
    return item

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(status_code=400, detail="Failed to extract text from DOCX file")

async def parse_document_sections(content: str, filename: str) -> Dict[str, str]:
    """Parse document content into logical sections"""
    sections = {}
    
    # Split by common section patterns
    patterns = [
        r'\b\d+\.\s+[A-Z][^.]*\n',  # Numbered sections
        r'\b[A-Z][^.]{10,50}:?\n',   # Header patterns
        r'^[A-Z\s]{3,}\n',           # All caps headers
    ]
    
    current_section = "Introduction"
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line is a section header
        is_header = False
        for pattern in patterns:
            if re.match(pattern, line + '\n'):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = line.replace(':', '').strip()
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(line)
    
    # Save final section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    # If no sections found, treat entire document as one section
    if not sections:
        sections[f"{filename} - Full Content"] = content
    
    return sections

async def analyze_clause_mapping(qsp_content: str, qsp_filename: str, iso_clauses: List[str]) -> List[Dict[str, Any]]:
    """Use AI to map QSP content to ISO clauses"""
    try:
        # Initialize LLM chat
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=f"mapping_{uuid.uuid4()}",
            system_message="""You are an expert in ISO 13485 compliance for medical devices. 
            Analyze QSP document sections and map them to relevant ISO 13485:2024 clauses.
            
            For each mapping, provide:
            1. The ISO clause number and title
            2. Confidence score (0.0-1.0) based on relevance
            3. Evidence text from the QSP that supports the mapping
            4. Brief explanation of the mapping
            
            Return results as a JSON array with this structure:
            [
                {
                    "iso_clause": "4.1 General requirements",
                    "confidence_score": 0.85,
                    "evidence_text": "relevant excerpt from QSP",
                    "explanation": "brief explanation of why this maps"
                }
            ]"""
        ).with_model("openai", "gpt-4o")
        
        iso_clauses_text = "\n".join([f"- {clause}" for clause in iso_clauses])
        
        user_message = UserMessage(
            text=f"""Analyze this QSP document content and map it to relevant ISO 13485:2024 clauses:

QSP Document: {qsp_filename}
Content:
{qsp_content[:3000]}...

ISO 13485:2024 Key Clauses to consider:
{iso_clauses_text}

Provide detailed mappings with confidence scores."""
        )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            # Clean the response - remove markdown formatting if present
            clean_response = response.strip()
            if clean_response.startswith('```json'):
                clean_response = clean_response[7:]  # Remove ```json
            if clean_response.endswith('```'):
                clean_response = clean_response[:-3]  # Remove ```
            clean_response = clean_response.strip()
            
            mappings = json.loads(clean_response)
            return mappings if isinstance(mappings, list) else []
        except json.JSONDecodeError:
            logger.error(f"Failed to parse AI response as JSON: {response}")
            return []
            
    except Exception as e:
        logger.error(f"Error in AI clause mapping: {e}")
        return []

# API Routes
@api_router.get("/")
async def root():
    return {"message": "QSP Compliance Checker API", "version": "1.0.0"}

# Health check endpoint (as requested in review)
@api_router.get("/health")
async def health_check():
    """Health check endpoint for database and AI services"""
    try:
        # Test database connection
        db_status = "healthy"
        try:
            await db.command("ping")
        except Exception as e:
            db_status = f"unhealthy: {str(e)}"
        
        # Test AI service
        ai_status = "healthy"
        try:
            chat = LlmChat(
                api_key=os.environ.get('EMERGENT_LLM_KEY'),
                session_id=f"health_{uuid.uuid4()}",
                system_message="You are a health check assistant."
            ).with_model("openai", "gpt-4o")
            
            user_message = UserMessage(text="Respond with 'OK' if you're working.")
            response = await chat.send_message(user_message)
            if "OK" not in response:
                ai_status = "unhealthy: unexpected response"
        except Exception as e:
            ai_status = f"unhealthy: {str(e)}"
        
        return {
            "status": "healthy" if db_status == "healthy" and ai_status == "healthy" else "degraded",
            "database": db_status,
            "ai_service": ai_status,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now(timezone.utc).isoformat()
        }

# Test endpoints (as requested in review)
@api_router.get("/test/database")
async def test_database():
    """Test MongoDB connectivity and operations"""
    try:
        # Test basic connection
        await db.command("ping")
        
        # Test collection operations
        test_doc = {"test": True, "timestamp": datetime.now(timezone.utc)}
        result = await db.test_collection.insert_one(test_doc)
        
        # Clean up test document
        await db.test_collection.delete_one({"_id": result.inserted_id})
        
        return {
            "status": "healthy",
            "message": "MongoDB connectivity and operations successful",
            "operations_tested": ["ping", "insert", "delete"]
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e)
        }

@api_router.get("/test/ai")
async def test_ai_service():
    """Test LLM service functionality with ISO 13485 explanation"""
    try:
        chat = LlmChat(
            api_key=os.environ.get('EMERGENT_LLM_KEY'),
            session_id=f"test_{uuid.uuid4()}",
            system_message="You are an expert in ISO 13485 medical device quality management."
        ).with_model("openai", "gpt-4o")
        
        user_message = UserMessage(
            text="Briefly explain what ISO 13485:2024 is and why it's important for medical device companies."
        )
        
        response = await chat.send_message(user_message)
        
        return {
            "status": "healthy",
            "message": "AI service is functioning correctly",
            "response": response,
            "model": "gpt-4o"
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e)
        }

@api_router.post("/test/upload")
async def test_document_upload(filename: str, content: str):
    """Simple test document upload endpoint"""
    try:
        # Create a simple test document
        test_doc = {
            "id": str(uuid.uuid4()),
            "filename": filename,
            "content": content,
            "upload_date": datetime.now(timezone.utc),
            "test_document": True
        }
        
        # Store in test collection
        await db.test_documents.insert_one(prepare_for_mongo(test_doc))
        
        return {
            "status": "success",
            "message": "Test document uploaded successfully",
            "document_id": test_doc["id"],
            "filename": filename,
            "content_length": len(content)
        }
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }

@api_router.get("/test/documents")
async def get_test_documents():
    """Retrieve uploaded test documents"""
    try:
        docs = await db.test_documents.find({}, {"_id": 0}).to_list(length=None)
        return [parse_from_mongo(doc) for doc in docs]
    except Exception as e:
        return {
            "status": "error",
            "error": str(e)
        }

@api_router.post("/documents/upload")
async def upload_qsp_document(file: UploadFile = File(...)):
    """Upload a QSP document (.docx or .txt)"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Check file type
        allowed_extensions = ['.docx', '.txt']
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,  # Changed from 500 to 400 for proper error code
                detail=f"Unsupported file type. Allowed: {allowed_extensions}"
            )
        
        # Read file content
        content = await file.read()
        
        # Extract text
        if file_ext == '.docx':
            text_content = await extract_text_from_docx(content)
        else:  # .txt
            text_content = content.decode('utf-8')
        
        # Parse into sections
        sections = await parse_document_sections(text_content, file.filename)
        
        # Create QSP document
        qsp_doc = QSPDocument(
            filename=file.filename,
            content=text_content,
            sections=sections,
            processed=True
        )
        
        # Store in MongoDB
        doc_dict = prepare_for_mongo(qsp_doc.model_dump())
        result = await db.qsp_documents.insert_one(doc_dict)
        
        logger.info(f"Uploaded QSP document: {file.filename}")
        
        return {
            "message": "Document uploaded successfully",
            "document_id": qsp_doc.id,
            "filename": file.filename,
            "sections_count": len(sections),
            "content_length": len(text_content)
        }
        
    except HTTPException:
        raise  # Re-raise HTTPExceptions to preserve status codes
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/iso-summary/upload")
async def upload_iso_summary(file: UploadFile = File(...)):
    """Upload ISO 13485:2024 Summary of Changes"""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        file_ext = Path(file.filename).suffix.lower()
        if file_ext == '.docx':
            text_content = await extract_text_from_docx(content)
        else:
            text_content = content.decode('utf-8')
        
        # Parse ISO summary content
        new_clauses = []
        modified_clauses = []
        
        # Simple parsing for new and modified clauses
        lines = text_content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if 'NEW CLAUSES' in line.upper() or 'NEW REQUIREMENTS' in line.upper():
                current_section = "new"
            elif 'MODIFIED CLAUSES' in line.upper() or 'CHANGED CLAUSES' in line.upper():
                current_section = "modified"
            elif line and current_section and re.match(r'\d+\.\d+', line):
                clause_info = {"clause": line, "description": line}
                if current_section == "new":
                    new_clauses.append(clause_info)
                elif current_section == "modified":
                    modified_clauses.append(clause_info)
        
        # Create ISO summary
        iso_summary = ISOSummary(
            content=text_content,
            new_clauses=new_clauses,
            modified_clauses=modified_clauses
        )
        
        # Store in MongoDB
        summary_dict = prepare_for_mongo(iso_summary.model_dump())
        await db.iso_summaries.insert_one(summary_dict)
        
        logger.info(f"Uploaded ISO summary: {file.filename}")
        
        return {
            "message": "ISO summary uploaded successfully",
            "summary_id": iso_summary.id,
            "new_clauses_count": len(new_clauses),
            "modified_clauses_count": len(modified_clauses)
        }
        
    except Exception as e:
        logger.error(f"Error uploading ISO summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/analysis/run-mapping")
async def run_clause_mapping():
    """Run AI-powered clause mapping for all QSP documents"""
    try:
        # Get all QSP documents
        qsp_docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        
        if not qsp_docs:
            raise HTTPException(status_code=404, detail="No QSP documents found")
        
        # Get ISO summary for clauses
        iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
        if not iso_summary:
            raise HTTPException(status_code=404, detail="No ISO summary found")
        
        # Standard ISO 13485 clauses
        iso_clauses = [
            "4.1 General requirements",
            "4.2 Documentation requirements",
            "5.1 Management commitment",
            "5.2 Customer focus",
            "5.3 Quality policy",
            "5.4 Planning",
            "5.5 Responsibility, authority and communication",
            "5.6 Management review",
            "6.1 Provision of resources",
            "6.2 Human resources",
            "6.3 Infrastructure",
            "6.4 Work environment",
            "7.1 Planning of product realization",
            "7.2 Customer-related processes",
            "7.3 Design and development",
            "7.4 Purchasing",
            "7.5 Production and service provision",
            "7.6 Control of monitoring and measuring equipment",
            "8.1 General",
            "8.2 Monitoring and measurement",
            "8.3 Control of nonconforming product",
            "8.4 Analysis of data",
            "8.5 Improvement"
        ]
        
        # Clear existing mappings
        await db.clause_mappings.delete_many({})
        
        total_mappings = 0
        
        # Process each QSP document
        for qsp_doc in qsp_docs:
            qsp_doc = parse_from_mongo(qsp_doc)
            
            # Map each section
            for section_title, section_content in qsp_doc['sections'].items():
                if len(section_content) < 50:  # Skip very short sections
                    continue
                
                mappings = await analyze_clause_mapping(
                    section_content, qsp_doc['filename'], iso_clauses
                )
                
                for mapping in mappings:
                    if mapping.get('confidence_score', 0) > 0.3:  # Only store high-confidence mappings
                        clause_mapping = ClauseMapping(
                            qsp_id=qsp_doc['id'],
                            qsp_filename=qsp_doc['filename'],
                            section_title=section_title,
                            section_content=section_content[:500],  # Truncate for storage
                            iso_clause=mapping['iso_clause'],
                            confidence_score=mapping['confidence_score'],
                            evidence_text=mapping['evidence_text']
                        )
                        
                        mapping_dict = prepare_for_mongo(clause_mapping.model_dump())
                        await db.clause_mappings.insert_one(mapping_dict)
                        total_mappings += 1
        
        logger.info(f"Generated {total_mappings} clause mappings")
        
        return {
            "message": "Clause mapping completed successfully",
            "total_documents_processed": len(qsp_docs),
            "total_mappings_generated": total_mappings
        }
        
    except Exception as e:
        logger.error(f"Error running clause mapping: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/analysis/run-compliance")
async def run_compliance_analysis():
    """Run compliance gap analysis"""
    try:
        # Get ISO summary and mappings
        iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
        if not iso_summary:
            raise HTTPException(status_code=404, detail="No ISO summary found")
        
        mappings = await db.clause_mappings.find({}, {"_id": 0}).to_list(length=None)
        qsp_docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        
        if not mappings:
            raise HTTPException(status_code=404, detail="No clause mappings found. Run mapping first.")
        
        # Extract changed clauses from ISO summary
        changed_clauses = set()
        iso_summary = parse_from_mongo(iso_summary)
        
        for clause in iso_summary.get('new_clauses', []):
            changed_clauses.add(clause.get('clause', ''))
        for clause in iso_summary.get('modified_clauses', []):
            changed_clauses.add(clause.get('clause', ''))
        
        # Find gaps
        gaps = []
        mapped_clauses = {}
        
        # Group mappings by clause
        for mapping in mappings:
            clause = mapping['iso_clause']
            if clause not in mapped_clauses:
                mapped_clauses[clause] = []
            mapped_clauses[clause].append(mapping)
        
        # Check for gaps in changed clauses
        for changed_clause in changed_clauses:
            if not changed_clause:
                continue
                
            # Find if any existing mappings cover this clause
            found_mapping = False
            low_confidence = []
            
            for clause, clause_mappings in mapped_clauses.items():
                if changed_clause in clause or any(changed_clause in m['iso_clause'] for m in clause_mappings):
                    found_mapping = True
                    # Check confidence scores
                    for mapping in clause_mappings:
                        if mapping['confidence_score'] < 0.7:
                            low_confidence.append(mapping)
            
            if not found_mapping:
                gap = ComplianceGap(
                    qsp_id="",
                    qsp_filename="Multiple",
                    iso_clause=changed_clause,
                    gap_type="missing",
                    description=f"No QSP documents found addressing the new/modified clause: {changed_clause}",
                    severity="high",
                    recommendations=[
                        f"Create or update QSP documentation to address {changed_clause}",
                        "Review existing procedures for compliance gaps",
                        "Assign responsibility for implementing new requirements"
                    ]
                )
                gaps.append(gap)
            elif low_confidence:
                for mapping in low_confidence:
                    gap = ComplianceGap(
                        qsp_id=mapping['qsp_id'],
                        qsp_filename=mapping['qsp_filename'],
                        iso_clause=mapping['iso_clause'],
                        gap_type="partial",
                        description=f"Low confidence mapping for {mapping['iso_clause']} in {mapping['qsp_filename']}",
                        severity="medium",
                        recommendations=[
                            "Review and strengthen documentation",
                            "Add more specific requirements",
                            "Clarify procedures and responsibilities"
                        ]
                    )
                    gaps.append(gap)
        
        # Calculate compliance score
        total_changed_clauses = len([c for c in changed_clauses if c])
        high_confidence_mappings = sum(1 for mapping in mappings if mapping['confidence_score'] > 0.7)
        
        if total_changed_clauses > 0:
            overall_score = min(high_confidence_mappings / total_changed_clauses * 100, 100)
        else:
            overall_score = 100
        
        # Get affected documents
        affected_docs = list(set(gap.qsp_filename for gap in gaps if gap.qsp_filename != "Multiple"))
        
        # Create analysis
        analysis = ComplianceAnalysis(
            overall_score=round(overall_score, 2),
            total_documents=len(qsp_docs),
            compliant_documents=len(qsp_docs) - len(affected_docs),
            gaps=gaps,
            affected_documents=affected_docs
        )
        
        # Store analysis
        analysis_dict = prepare_for_mongo(analysis.model_dump())
        await db.compliance_analyses.delete_many({})  # Keep only latest
        await db.compliance_analyses.insert_one(analysis_dict)
        
        logger.info(f"Compliance analysis completed. Found {len(gaps)} gaps.")
        
        return {
            "message": "Compliance analysis completed successfully",
            "overall_score": analysis.overall_score,
            "gaps_found": len(gaps),
            "affected_documents": len(affected_docs),
            "analysis_id": analysis.id
        }
        
    except Exception as e:
        logger.error(f"Error running compliance analysis: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/dashboard")
async def get_dashboard_data():
    """Get dashboard overview data"""
    try:
        # Get latest analysis
        analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
        
        # Get document counts
        total_docs = await db.qsp_documents.count_documents({})
        total_mappings = await db.clause_mappings.count_documents({})
        
        # Get ISO summary
        iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
        
        if analysis:
            analysis = parse_from_mongo(analysis)
        
        if iso_summary:
            iso_summary = parse_from_mongo(iso_summary)
        
        return {
            "compliance_score": analysis['overall_score'] if analysis else 0,
            "total_documents": total_docs,
            "total_mappings": total_mappings,
            "gaps_count": len(analysis['gaps']) if analysis else 0,
            "affected_documents": len(analysis['affected_documents']) if analysis else 0,
            "iso_summary_loaded": iso_summary is not None,
            "new_clauses_count": len(iso_summary.get('new_clauses', [])) if iso_summary else 0,
            "modified_clauses_count": len(iso_summary.get('modified_clauses', [])) if iso_summary else 0,
            "last_analysis_date": analysis['analysis_date'] if analysis else None
        }
        
    except Exception as e:
        logger.error(f"Error getting dashboard data: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents", response_model=List[Dict[str, Any]])
async def get_qsp_documents():
    """Get all QSP documents"""
    try:
        docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        return [parse_from_mongo(doc) for doc in docs]
    except Exception as e:
        logger.error(f"Error getting documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/gaps")
async def get_compliance_gaps():
    """Get detailed compliance gaps"""
    try:
        analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
        
        if not analysis:
            return {"gaps": [], "message": "No compliance analysis found. Run analysis first."}
        
        analysis = parse_from_mongo(analysis)
        
        return {
            "gaps": analysis['gaps'],
            "total_gaps": len(analysis['gaps']),
            "analysis_date": analysis['analysis_date']
        }
        
    except Exception as e:
        logger.error(f"Error getting compliance gaps: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/mappings")
async def get_clause_mappings():
    """Get all clause mappings"""
    try:
        mappings = await db.clause_mappings.find({}, {"_id": 0}).to_list(length=None)
        return [parse_from_mongo(mapping) for mapping in mappings]
    except Exception as e:
        logger.error(f"Error getting mappings: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Include the router in the main app
app.include_router(api_router)

# Store active WebSocket connections for MCP
active_mcp_connections: Dict[str, WebSocket] = {}

# MCP WebSocket endpoint
@app.websocket("/mcp")
async def mcp_websocket_endpoint(websocket: WebSocket):
    """WebSocket endpoint for MCP protocol communication with Claude/ChatGPT"""
    await websocket.accept()
    connection_id = str(uuid.uuid4())
    active_mcp_connections[connection_id] = websocket
    
    logger.info(f"New MCP connection: {connection_id}")
    
    try:
        while True:
            # Receive message from client (Claude/ChatGPT)
            data = await websocket.receive_text()
            message = json.loads(data)
            
            logger.info(f"Received MCP message: {message.get('method', 'unknown')}")
            
            # Handle MCP protocol messages
            response = await handle_mcp_protocol_message(message)
            
            # Send response back
            await websocket.send_text(json.dumps(response))
            
    except Exception as e:
        logger.error(f"MCP WebSocket error: {e}")
    finally:
        if connection_id in active_mcp_connections:
            del active_mcp_connections[connection_id]
        logger.info(f"MCP connection closed: {connection_id}")

async def handle_mcp_protocol_message(message: Dict[str, Any]) -> Dict[str, Any]:
    """Handle incoming MCP protocol messages"""
    
    method = message.get("method")
    params = message.get("params", {})
    message_id = message.get("id")
    
    try:
        if method == "initialize":
            return {
                "jsonrpc": "2.0",
                "id": message_id,
                "result": {
                    "protocolVersion": "2024-11-05",
                    "capabilities": {
                        "tools": {},
                        "prompts": {},
                        "resources": {}
                    },
                    "serverInfo": {
                        "name": "qsp-compliance-checker",
                        "version": "1.0.0"
                    }
                }
            }
        
        elif method == "tools/list":
            tools = [
                {
                    "name": "upload_qsp_document",
                    "description": "Upload a QSP document for compliance analysis",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "string"},
                            "content": {"type": "string", "description": "Base64 encoded file content"},
                            "file_type": {"type": "string", "enum": ["txt", "docx"]}
                        },
                        "required": ["filename", "content", "file_type"]
                    }
                },
                {
                    "name": "upload_iso_summary", 
                    "description": "Upload ISO 13485:2024 Summary of Changes",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "filename": {"type": "string"},
                            "content": {"type": "string", "description": "Base64 encoded file content"},
                            "file_type": {"type": "string", "enum": ["txt", "docx", "pdf"]}
                        },
                        "required": ["filename", "content", "file_type"]
                    }
                },
                {
                    "name": "list_documents",
                    "description": "List all uploaded QSP documents",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                        }
                    }
                },
                {
                    "name": "run_clause_mapping",
                    "description": "Run AI-powered clause mapping analysis",
                    "inputSchema": {"type": "object", "properties": {}}
                },
                {
                    "name": "run_compliance_analysis", 
                    "description": "Run comprehensive compliance gap analysis",
                    "inputSchema": {"type": "object", "properties": {}}
                },
                {
                    "name": "get_compliance_status",
                    "description": "Get current compliance status and metrics",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                        }
                    }
                },
                {
                    "name": "get_dashboard_summary",
                    "description": "Get dashboard overview with key metrics",
                    "inputSchema": {
                        "type": "object", 
                        "properties": {
                            "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                        }
                    }
                },
                {
                    "name": "get_detailed_gaps",
                    "description": "Get detailed compliance gaps with recommendations",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "severity": {"type": "string", "enum": ["all", "high", "medium", "low"], "default": "all"},
                            "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                        }
                    }
                },
                {
                    "name": "get_clause_mappings",
                    "description": "Get AI clause mappings between QSP and ISO clauses",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "document_id": {"type": "string", "description": "Filter by document ID"},
                            "iso_clause": {"type": "string", "description": "Filter by ISO clause"},
                            "min_confidence": {"type": "number", "minimum": 0, "maximum": 1},
                            "format": {"type": "string", "enum": ["summary", "detailed", "json"], "default": "summary"}
                        }
                    }
                },
                {
                    "name": "query_specific_clause",
                    "description": "Query compliance for specific ISO clause",
                    "inputSchema": {
                        "type": "object",
                        "properties": {
                            "clause": {"type": "string", "description": "ISO clause to query"},
                            "include_recommendations": {"type": "boolean", "default": True}
                        },
                        "required": ["clause"]
                    }
                }
            ]
            
            return {
                "jsonrpc": "2.0",
                "id": message_id,
                "result": {"tools": tools}
            }
        
        elif method == "tools/call":
            tool_name = params.get("name")
            arguments = params.get("arguments", {})
            
            # Convert base64 content for file uploads
            if tool_name in ["upload_qsp_document", "upload_iso_summary"] and "content" in arguments:
                arguments["base64_content"] = arguments.pop("content", "")
            
            # Call the appropriate tool
            result_text = await call_mcp_tool(tool_name, arguments)
            
            return {
                "jsonrpc": "2.0", 
                "id": message_id,
                "result": {
                    "content": [{"type": "text", "text": result_text}]
                }
            }
        
        else:
            return {
                "jsonrpc": "2.0",
                "id": message_id,
                "error": {
                    "code": -32601,
                    "message": f"Method not found: {method}"
                }
            }
    
    except Exception as e:
        logger.error(f"Error handling MCP message: {e}")
        return {
            "jsonrpc": "2.0",
            "id": message_id, 
            "error": {
                "code": -32603,
                "message": f"Internal error: {str(e)}"
            }
        }

async def call_mcp_tool(tool_name: str, arguments: dict) -> str:
    """Call MCP tool and return formatted response"""
    try:
        if tool_name == "upload_qsp_document":
            return await mcp_upload_qsp_document(arguments)
        elif tool_name == "upload_iso_summary":
            return await mcp_upload_iso_summary(arguments)
        elif tool_name == "list_documents":
            return await mcp_list_documents(arguments)
        elif tool_name == "run_clause_mapping":
            return await mcp_run_clause_mapping(arguments)
        elif tool_name == "run_compliance_analysis":
            return await mcp_run_compliance_analysis(arguments)
        elif tool_name == "get_compliance_status":
            return await mcp_get_compliance_status(arguments)
        elif tool_name == "get_dashboard_summary":
            return await mcp_get_compliance_status(arguments)  # Same as status
        elif tool_name == "get_detailed_gaps":
            return await mcp_get_detailed_gaps(arguments)
        elif tool_name == "get_clause_mappings":
            return await mcp_get_clause_mappings(arguments)
        elif tool_name == "query_specific_clause":
            return await mcp_query_specific_clause(arguments)
        else:
            return f"❌ Unknown tool: {tool_name}"
    except Exception as e:
        logger.error(f"Error executing MCP tool {tool_name}: {e}")
        return f"❌ Error executing {tool_name}: {str(e)}"

# MCP Tool Implementations (simplified versions of existing functions)
async def mcp_upload_qsp_document(args: dict) -> str:
    """MCP version of document upload"""
    filename = args["filename"]
    file_type = args["file_type"]
    base64_content = args["base64_content"]
    
    try:
        # Decode base64 content
        content = base64.b64decode(base64_content)
        
        # Extract text
        if file_type == "docx":
            doc = Document(io.BytesIO(content))
            text_content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:  # txt
            text_content = content.decode('utf-8')
        
        # Parse into sections (simplified)
        sections = {}
        lines = text_content.split('\n')
        current_section = "Introduction"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if any(pattern in line.upper() for pattern in ['PURPOSE', 'SCOPE', 'PROCEDURE', 'RESPONSIBILITIES']):
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        if not sections:
            sections[f"{filename} - Full Content"] = text_content
        
        # Create QSP document
        qsp_doc = QSPDocument(
            filename=filename,
            content=text_content,
            sections=sections,
            processed=True
        )
        
        # Store in MongoDB
        doc_dict = prepare_for_mongo(qsp_doc.model_dump())
        await db.qsp_documents.insert_one(doc_dict)
        
        return f"✅ **QSP Document Uploaded Successfully**\n\n**File:** {filename}\n**Sections:** {len(sections)}\n**Content:** {len(text_content)} characters\n\nReady for clause mapping analysis!"
        
    except Exception as e:
        return f"❌ Upload failed: {str(e)}"

async def mcp_upload_iso_summary(args: dict) -> str:
    """MCP version of ISO summary upload"""
    filename = args["filename"] 
    file_type = args["file_type"]
    base64_content = args["base64_content"]
    
    try:
        # Decode base64 content
        content = base64.b64decode(base64_content)
        
        # Extract text
        if file_type == "docx":
            doc = Document(io.BytesIO(content))
            text_content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:  # txt
            text_content = content.decode('utf-8')
        
        # Parse clauses (simplified)
        new_clauses = []
        modified_clauses = []
        
        lines = text_content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if 'NEW CLAUSES' in line.upper():
                current_section = "new"
            elif 'MODIFIED CLAUSES' in line.upper():
                current_section = "modified"
            elif line and current_section and re.match(r'\d+\.\d+', line):
                clause_info = {"clause": line, "description": line}
                if current_section == "new":
                    new_clauses.append(clause_info)
                elif current_section == "modified":
                    modified_clauses.append(clause_info)
        
        # Create ISO summary
        iso_summary = ISOSummary(
            content=text_content,
            new_clauses=new_clauses,
            modified_clauses=modified_clauses
        )
        
        # Store in MongoDB
        summary_dict = prepare_for_mongo(iso_summary.model_dump())
        await db.iso_summaries.insert_one(summary_dict)
        
        return f"✅ **ISO 13485:2024 Summary Uploaded**\n\n**File:** {filename}\n**New Clauses:** {len(new_clauses)}\n**Modified Clauses:** {len(modified_clauses)}\n\nCompliance analysis enabled!"
        
    except Exception as e:
        return f"❌ Upload failed: {str(e)}"

async def mcp_list_documents(args: dict) -> str:
    """MCP version of list documents"""
    try:
        docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        
        if not docs:
            return "📄 **No QSP Documents Found**\n\nUpload documents to get started with compliance analysis."
        
        result = f"📄 **{len(docs)} QSP Documents**\n\n"
        for doc in docs:
            doc = parse_from_mongo(doc)
            result += f"• **{doc['filename']}** - {len(doc.get('sections', {}))} sections\n"
            result += f"  Uploaded: {doc['upload_date'].strftime('%Y-%m-%d %H:%M')}\n\n"
        
        return result
        
    except Exception as e:
        return f"❌ Error listing documents: {str(e)}"

async def mcp_run_clause_mapping(args: dict) -> str:
    """MCP version of clause mapping"""
    try:
        qsp_docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        if not qsp_docs:
            return "❌ No QSP documents found. Upload documents first."
        
        # Clear existing mappings
        await db.clause_mappings.delete_many({})
        
        # Standard ISO clauses
        iso_clauses = [
            "4.1 General requirements", "4.2 Documentation requirements",
            "5.1 Management commitment", "7.3 Design and development",
            "7.4 Purchasing", "8.1 General", "8.5 Improvement"
        ]
        
        total_mappings = 0
        
        # Process each document
        for qsp_doc in qsp_docs:
            qsp_doc = parse_from_mongo(qsp_doc)
            
            for section_title, section_content in qsp_doc['sections'].items():
                if len(section_content) < 50:
                    continue
                
                mappings = await analyze_clause_mapping(
                    section_content, qsp_doc['filename'], iso_clauses
                )
                
                for mapping in mappings:
                    if mapping.get('confidence_score', 0) > 0.3:
                        clause_mapping = ClauseMapping(
                            qsp_id=qsp_doc['id'],
                            qsp_filename=qsp_doc['filename'],
                            section_title=section_title,
                            section_content=section_content[:500],
                            iso_clause=mapping['iso_clause'],
                            confidence_score=mapping['confidence_score'],
                            evidence_text=mapping['evidence_text']
                        )
                        
                        mapping_dict = prepare_for_mongo(clause_mapping.model_dump())
                        await db.clause_mappings.insert_one(mapping_dict)
                        total_mappings += 1
        
        return f"🤖 **AI Clause Mapping Complete**\n\n**Documents Processed:** {len(qsp_docs)}\n**Mappings Generated:** {total_mappings}\n\nYour QSP documents have been analyzed and mapped to ISO clauses. Run compliance analysis next!"
        
    except Exception as e:
        return f"❌ Mapping failed: {str(e)}"

async def mcp_run_compliance_analysis(args: dict) -> str:
    """MCP version of compliance analysis"""
    try:
        # Get data
        iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
        if not iso_summary:
            return "❌ No ISO summary found. Upload ISO 13485:2024 Summary first."
        
        mappings = await db.clause_mappings.find({}, {"_id": 0}).to_list(length=None)
        if not mappings:
            return "❌ No clause mappings found. Run clause mapping first."
        
        qsp_docs = await db.qsp_documents.find({}, {"_id": 0}).to_list(length=None)
        
        # Extract changed clauses
        changed_clauses = set()
        iso_summary = parse_from_mongo(iso_summary)
        
        for clause in iso_summary.get('new_clauses', []):
            changed_clauses.add(clause.get('clause', ''))
        for clause in iso_summary.get('modified_clauses', []):
            changed_clauses.add(clause.get('clause', ''))
        
        # Find gaps
        gaps = []
        mapped_clauses = {mapping['iso_clause'] for mapping in mappings}
        
        for changed_clause in changed_clauses:
            if not changed_clause:
                continue
            
            found_mapping = any(changed_clause in clause for clause in mapped_clauses)
            
            if not found_mapping:
                gap = ComplianceGap(
                    qsp_id="",
                    qsp_filename="Multiple",
                    iso_clause=changed_clause,
                    gap_type="missing",
                    description=f"No QSP documents address the clause: {changed_clause}",
                    severity="high",
                    recommendations=[
                        f"Create documentation for {changed_clause}",
                        "Review existing procedures for gaps",
                        "Assign implementation responsibility"
                    ]
                )
                gaps.append(gap)
        
        # Calculate score
        total_changed = len([c for c in changed_clauses if c])
        high_conf_mappings = sum(1 for m in mappings if m['confidence_score'] > 0.7)
        overall_score = min(high_conf_mappings / max(total_changed, 1) * 100, 100)
        
        # Create analysis
        analysis = ComplianceAnalysis(
            overall_score=round(overall_score, 2),
            total_documents=len(qsp_docs),
            compliant_documents=len(qsp_docs) - len(set(g.qsp_filename for g in gaps if g.qsp_filename != "Multiple")),
            gaps=gaps,
            affected_documents=[]
        )
        
        # Store analysis
        analysis_dict = prepare_for_mongo(analysis.model_dump())
        await db.compliance_analyses.delete_many({})
        await db.compliance_analyses.insert_one(analysis_dict)
        
        high_gaps = len([g for g in gaps if g.severity == "high"])
        
        return f"📊 **Compliance Analysis Complete**\n\n**Overall Score:** {analysis.overall_score}%\n**Documents Analyzed:** {analysis.total_documents}\n**High Priority Gaps:** {high_gaps}\n\nAnalysis complete! Check detailed gaps for specific recommendations."
        
    except Exception as e:
        return f"❌ Analysis failed: {str(e)}"

async def mcp_get_compliance_status(args: dict) -> str:
    """MCP version of compliance status"""
    try:
        # Get latest analysis
        analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
        total_docs = await db.qsp_documents.count_documents({})
        total_mappings = await db.clause_mappings.count_documents({})
        iso_summary = await db.iso_summaries.find_one({}, {"_id": 0})
        
        if analysis:
            analysis = parse_from_mongo(analysis)
        
        score = analysis['overall_score'] if analysis else 0
        gaps = len(analysis['gaps']) if analysis else 0
        
        status = f"📊 **QSP Compliance Status**\n\n"
        status += f"**Overall Compliance Score:** {score}%\n"
        status += f"**Documents Analyzed:** {total_docs}\n" 
        status += f"**Compliance Gaps:** {gaps} issues identified\n"
        status += f"**AI Clause Mappings:** {total_mappings} generated\n"
        status += f"**ISO Summary:** {'✅ Loaded' if iso_summary else '❌ Not uploaded'}\n\n"
        
        if analysis:
            status += f"Last Analysis: {analysis['analysis_date'].strftime('%Y-%m-%d %H:%M')}"
        else:
            status += "⚠️ No analysis run yet. Upload documents and run analysis."
        
        return status
        
    except Exception as e:
        return f"❌ Error getting status: {str(e)}"

async def mcp_get_detailed_gaps(args: dict) -> str:
    """MCP version of detailed gaps"""
    try:
        analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
        
        if not analysis:
            return "❌ No compliance analysis found. Run analysis first."
        
        analysis = parse_from_mongo(analysis)
        gaps = analysis['gaps']
        
        severity_filter = args.get("severity", "all")
        if severity_filter != "all":
            gaps = [gap for gap in gaps if gap['severity'] == severity_filter]
        
        if not gaps:
            return "✅ **No Compliance Gaps Found**\n\nExcellent! Your QSP documents appear to be compliant."
        
        result = f"⚠️ **{len(gaps)} Compliance Gaps**\n\n"
        
        for i, gap in enumerate(gaps[:5], 1):  # Show first 5
            result += f"**{i}. {gap['iso_clause']}** ({gap['severity'].upper()})\n"
            result += f"📋 Document: {gap['qsp_filename']}\n"
            result += f"📝 Issue: {gap['description']}\n"
            
            if gap['recommendations']:
                result += f"💡 Recommendations:\n"
                for rec in gap['recommendations'][:2]:
                    result += f"   • {rec}\n"
            result += "\n"
        
        if len(gaps) > 5:
            result += f"... and {len(gaps) - 5} more gaps"
        
        return result
        
    except Exception as e:
        return f"❌ Error getting gaps: {str(e)}"

async def mcp_get_clause_mappings(args: dict) -> str:
    """MCP version of clause mappings"""
    try:
        query = {}
        if args.get("document_id"):
            query["qsp_id"] = args["document_id"]
        if args.get("iso_clause"):
            query["iso_clause"] = {"$regex": args["iso_clause"], "$options": "i"}
        if args.get("min_confidence", 0) > 0:
            query["confidence_score"] = {"$gte": args["min_confidence"]}
        
        mappings = await db.clause_mappings.find(query, {"_id": 0}).to_list(length=None)
        
        if not mappings:
            return "🔗 **No Clause Mappings Found**\n\nRun clause mapping analysis first."
        
        result = f"🔗 **{len(mappings)} Clause Mappings**\n\n"
        
        for mapping in mappings[:8]:  # Show first 8
            result += f"• **{mapping['iso_clause']}** (confidence: {mapping['confidence_score']:.2f})\n"
            result += f"  📄 {mapping['qsp_filename']} - {mapping['section_title']}\n"
            result += f"  📝 Evidence: {mapping['evidence_text'][:80]}...\n\n"
        
        if len(mappings) > 8:
            result += f"... and {len(mappings) - 8} more mappings"
        
        return result
        
    except Exception as e:
        return f"❌ Error getting mappings: {str(e)}"

async def mcp_query_specific_clause(args: dict) -> str:
    """MCP version of specific clause query"""
    try:
        clause = args["clause"]
        
        # Find mappings
        mappings = await db.clause_mappings.find(
            {"iso_clause": {"$regex": clause, "$options": "i"}},
            {"_id": 0}
        ).to_list(length=None)
        
        # Find gaps
        analysis = await db.compliance_analyses.find_one({}, {"_id": 0}, sort=[("analysis_date", -1)])
        relevant_gaps = []
        
        if analysis:
            analysis = parse_from_mongo(analysis)
            relevant_gaps = [gap for gap in analysis['gaps'] if clause in gap['iso_clause']]
        
        result = f"🔍 **Analysis for ISO Clause {clause}**\n\n"
        
        if mappings:
            result += f"**📋 QSP Coverage ({len(mappings)} mappings):**\n"
            for mapping in mappings[:3]:
                result += f"• **{mapping['qsp_filename']}** - {mapping['section_title']}\n"
                result += f"  Confidence: {mapping['confidence_score']:.2f}\n"
                result += f"  Evidence: {mapping['evidence_text'][:60]}...\n\n"
        else:
            result += "**📋 No QSP Coverage Found**\n\n"
        
        if relevant_gaps:
            result += f"**⚠️ Compliance Gaps ({len(relevant_gaps)}):**\n"
            for gap in relevant_gaps:
                result += f"• **{gap['severity'].upper()}:** {gap['description']}\n"
                if args.get("include_recommendations", True) and gap['recommendations']:
                    result += f"  💡 Recommendation: {gap['recommendations'][0]}\n"
                result += "\n"
        else:
            result += "**✅ No compliance gaps found for this clause**\n"
        
        return result
        
    except Exception as e:
        return f"❌ Error querying clause: {str(e)}"

# MCP Setup endpoint
@api_router.get("/mcp/setup")
async def get_mcp_setup_instructions():
    """Get MCP setup instructions for Claude Desktop and ChatGPT"""
    domain = os.environ.get('DOMAIN', 'qsp-compliance.preview.emergentagent.com')
    
    return {
        "title": "QSP Compliance Checker - MCP Setup",
        "status": "ready",
        "websocket_endpoint": f"wss://{domain}/mcp",
        "claude_desktop_config": {
            "instructions": "Add this to your Claude Desktop configuration file:",
            "config": {
                "mcpServers": {
                    "qsp-compliance": {
                        "url": f"wss://{domain}/mcp",
                        "name": "QSP Compliance Checker"
                    }
                }
            },
            "config_file_locations": {
                "macOS": "~/Library/Application Support/Claude/claude_desktop_config.json",
                "Windows": "%APPDATA%/Claude/claude_desktop_config.json",
                "Linux": "~/.config/claude/claude_desktop_config.json"
            }
        },
        "chatgpt_config": {
            "instructions": "When ChatGPT supports MCP, use this WebSocket URL:",
            "url": f"wss://{domain}/mcp"
        },
        "test_commands": [
            "What QSP compliance tools do you have access to?",
            "Show me the current compliance status",
            "List all uploaded QSP documents",
            "Upload this QSP document: [paste document content]"
        ],
        "tools_available": 10,
        "features": [
            "Document upload via base64 encoding",
            "AI-powered clause mapping", 
            "Compliance gap analysis",
            "ISO 13485:2024 compliance checking",
            "Real-time status reporting"
        ]
    }

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()