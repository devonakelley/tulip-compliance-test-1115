"""
QSP Document Parser
Parses DOCX/PDF/TXT QSP files into structured clause-level data
Extracts document number, clause numbers, titles, and text content
"""
import logging
import re
from typing import List, Dict, Any, Optional
from docx import Document
from pathlib import Path
import io
from .reference_extractor import get_reference_extractor

logger = logging.getLogger(__name__)


class QSPParser:
    """
    Parse QSP documents into structured clause-level data
    """
    
    def __init__(self):
        self.heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5']
    
    def extract_document_number(self, filename: str) -> str:
        """
        Extract document number from filename
        Examples:
        - "QSP 7.3-3 R9 Risk Management.docx" -> "7.3-3"
        - "QSP 6.2-1 R16 Training and Qualification.docx" -> "6.2-1"
        - "QSP 4.2-1 R13 Document Control.docx" -> "4.2-1"
        """
        try:
            # Pattern to match: QSP X.X-X or just X.X-X
            pattern = r'(?:QSP\s+)?(\d+\.\d+-\d+)'
            match = re.search(pattern, filename)
            if match:
                return match.group(1)
            
            # Fallback: try simpler pattern X.X
            pattern2 = r'(?:QSP\s+)?(\d+\.\d+)'
            match2 = re.search(pattern2, filename)
            if match2:
                return match2.group(1)
            
            # If no pattern matched, return filename without extension
            return Path(filename).stem
            
        except Exception as e:
            logger.error(f"Failed to extract document number from {filename}: {e}")
            return Path(filename).stem
    
    def extract_revision(self, filename: str) -> str:
        """
        Extract revision from filename
        Examples:
        - "QSP 7.3-3 R9 Risk Management.docx" -> "R9"
        - "QSP 6.2-1 R16 Training.docx" -> "R16"
        """
        try:
            pattern = r'R(\d+)'
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                return f"R{match.group(1)}"
            return "Unknown"
        except Exception as e:
            logger.error(f"Failed to extract revision from {filename}: {e}")
            return "Unknown"
    
    def is_heading_paragraph(self, paragraph) -> bool:
        """Check if a paragraph is a heading using the exact specified pattern"""
        if not paragraph.text.strip():
            return False
        
        text = paragraph.text.strip()
        
        # Primary pattern from requirements: ^\d+(\.\d+)+\s+[A-Z][A-Za-z\s-]+
        # This matches: "4.2.1 Purpose", "7.3.5 Risk Analysis", etc.
        clause_pattern = r'^\d+(\.\d+)+\s+[A-Z][A-Za-z\s\-]+'
        if re.match(clause_pattern, text):
            return True
        
        return False
    
    def extract_clause_number(self, heading_text: str) -> Optional[str]:
        """
        Extract clause number from heading text
        Examples:
        - "7.3.5 Risk Analysis" -> "7.3.5"
        - "4.2.1 General Requirements" -> "4.2.1"
        - "Section 6.2 Training" -> "6.2"
        """
        try:
            # Pattern for X.X.X or X.X format at start of heading
            pattern = r'^(?:Section\s+)?(\d+(?:\.\d+)+)'
            match = re.search(pattern, heading_text.strip())
            if match:
                return match.group(1)
            return None
        except Exception as e:
            logger.error(f"Failed to extract clause number from '{heading_text}': {e}")
            return None
    
    def is_noise_line(self, text: str) -> bool:
        """
        Filter out noise lines per requirements:
        - Lines shorter than 50 characters
        - Lines containing: Date, Signature, Tulip Medical, Approvals
        """
        if not text:
            return True
        
        text_stripped = text.strip()
        
        # Skip lines shorter than 50 characters
        if len(text_stripped) < 50:
            return True
        
        # Skip lines with noise keywords (case-insensitive)
        noise_keywords = [
            'Date', 'Signature', 'Tulip Medical', 'Approvals',
            'Approval', 'Approved By', 'Reviewed By', 'Prepared By'
        ]
        
        text_lower = text_stripped.lower()
        for noise in noise_keywords:
            if noise.lower() in text_lower:
                return True
        
        return False
    
    def parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX file into structured clause-level data with proper text aggregation
        
        Returns:
        {
            "document_number": "7.3-3",
            "revision": "R9",
            "filename": "QSP 7.3-3 R9 Risk Management.docx",
            "clauses": [
                {
                    "document_number": "7.3-3",
                    "clause_number": "7.3.5",
                    "title": "Risk Analysis",
                    "text": "The risk analysis can be recorded on Form 7.3-3-2...",
                    "characters": 312
                }
            ]
        }
        """
        try:
            doc = Document(io.BytesIO(file_content))
            document_number = self.extract_document_number(filename)
            revision = self.extract_revision(filename)
            
            clauses = []
            current_heading = None
            current_clause_number = None
            current_text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # Check if this is a noise line (skip it)
                if self.is_noise_line(text):
                    continue
                
                # Check if this is a heading
                if self.is_heading_paragraph(paragraph):
                    # Save previous section if it has content
                    if current_heading and current_text_parts:
                        section_text = '\n'.join(current_text_parts).strip()
                        
                        # Only save if we have substantial text (not just a header)
                        if len(section_text) >= 50:
                            # Extract references from this clause's text
                            ref_extractor = get_reference_extractor()
                            references = ref_extractor.extract_all_references(section_text)
                            
                            clauses.append({
                                "document_number": document_number,
                                "revision": revision,
                                "clause": current_clause_number if current_clause_number else "Unknown",
                                "title": current_heading,
                                "characters": len(section_text),
                                "text": section_text,
                                "references": references  # NEW FIELD
                            })
                    
                    # Start new section
                    current_heading = text
                    current_clause_number = self.extract_clause_number(text)
                    
                    # If no clause number found but it's a common header, use semantic naming
                    if not current_clause_number:
                        common_headers = ['PURPOSE', 'SCOPE', 'RESPONSIBILITIES', 'PROCEDURE', 'RECORDS']
                        if text.upper() in common_headers:
                            current_clause_number = f"{document_number}.{text.upper()}"
                    
                    current_text_parts = []
                else:
                    # This is content - add to current section
                    if current_heading and len(text) > 20:  # Filter out very short lines
                        current_text_parts.append(text)
            
            # Save last section if it has content
            if current_heading and current_text_parts:
                section_text = '\n'.join(current_text_parts).strip()
                if len(section_text) >= 50:
                    # Extract references from this clause's text
                    ref_extractor = get_reference_extractor()
                    references = ref_extractor.extract_all_references(section_text)
                    
                    clauses.append({
                        "document_number": document_number,
                        "revision": revision,
                        "clause": current_clause_number if current_clause_number else "Unknown",
                        "title": current_heading,
                        "characters": len(section_text),
                        "text": section_text,
                        "references": references  # NEW FIELD
                    })
            
            # Fallback: if no clauses found, try alternative parsing
            if len(clauses) == 0:
                logger.warning(f"No clauses found with primary method, trying fallback for {filename}")
                clauses = self._fallback_parsing(doc, document_number)
            
            logger.info(f"✅ Parsed {filename}: {len(clauses)} clauses extracted")
            
            return {
                "document_number": document_number,
                "revision": revision,
                "filename": filename,
                "total_clauses": len(clauses),
                "clauses": clauses
            }
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {filename}: {e}")
            raise
    
    def _fallback_parsing(self, doc, document_number: str) -> List[Dict[str, Any]]:
        """
        Fallback parsing when primary method fails
        Looks for any structured content and creates minimal sections
        """
        clauses = []
        all_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and not self.is_noise_line(text):
                all_text.append(text)
        
        # If we have substantial content, create a single catch-all section
        if all_text:
            full_text = '\n'.join(all_text)
            if len(full_text) >= 100:
                clauses.append({
                    "document_number": document_number,
                    "clause_number": f"{document_number}.1",
                    "title": "Document Content",
                    "text": full_text,  # No truncation
                    "characters": len(full_text)
                })
        
        return clauses
    
    def parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF file into structured clause-level data
        Uses PyPDF2 for text extraction
        """
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract all text
            full_text = ""
            for page in pdf_reader.pages:
                full_text += page.extract_text() + "\n"
            
            # Parse text into sections (similar logic to DOCX)
            document_number = self.extract_document_number(filename)
            revision = self.extract_revision(filename)
            
            clauses = self._parse_text_into_sections(full_text, document_number)
            
            logger.info(f"✅ Parsed PDF {filename}: {len(clauses)} clauses extracted")
            
            return {
                "document_number": document_number,
                "revision": revision,
                "filename": filename,
                "total_clauses": len(clauses),
                "clauses": clauses
            }
            
        except Exception as e:
            logger.error(f"Failed to parse PDF file {filename}: {e}")
            raise
    
    def parse_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse TXT file into structured clause-level data
        """
        try:
            text = file_content.decode('utf-8', errors='ignore')
            document_number = self.extract_document_number(filename)
            revision = self.extract_revision(filename)
            
            clauses = self._parse_text_into_sections(text, document_number)
            
            logger.info(f"✅ Parsed TXT {filename}: {len(clauses)} clauses extracted")
            
            return {
                "document_number": document_number,
                "revision": revision,
                "filename": filename,
                "total_clauses": len(clauses),
                "clauses": clauses
            }
            
        except Exception as e:
            logger.error(f"Failed to parse TXT file {filename}: {e}")
            raise
    
    def _parse_text_into_sections(self, text: str, document_number: str) -> List[Dict[str, Any]]:
        """
        Parse plain text into sections based on heading patterns
        Used for PDF and TXT files with improved clause detection
        """
        clauses = []
        lines = text.split('\n')
        
        current_heading = None
        current_clause_number = None
        current_text_parts = []
        
        # Clause pattern: X.X.X Section Name
        clause_pattern = r'^[0-9]+(\.[0-9]+)+\s+[A-Z][A-Za-z\s\-]+'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip noise
            if self.is_noise_line(line):
                continue
            
            # Check if line is a heading
            is_heading = False
            
            # Pattern 1: Numbered clause (e.g., "4.2.1 Purpose")
            if re.match(clause_pattern, line):
                is_heading = True
            
            # Pattern 2: Common uppercase headers
            elif line.upper() in ['PURPOSE', 'SCOPE', 'RESPONSIBILITIES', 'PROCEDURE', 'RECORDS', 
                                  'DEFINITIONS', 'REFERENCES', 'ATTACHMENTS', 'FORMS', 'BACKGROUND']:
                is_heading = True
            
            # Pattern 3: Short uppercase line (but not noise)
            elif len(line) < 80 and line.isupper() and len(line) > 15:
                is_heading = True
            
            if is_heading:
                # Save previous section if it has content
                if current_heading and current_text_parts:
                    section_text = '\n'.join(current_text_parts).strip()
                    
                    if len(section_text) >= 50:
                        clauses.append({
                            "document_number": document_number,
                            "clause_number": current_clause_number if current_clause_number else f"{document_number}.{len(clauses)+1}",
                            "title": current_heading,
                            "text": section_text,
                            "characters": len(section_text)
                        })
                
                # Start new section
                current_heading = line
                current_clause_number = self.extract_clause_number(line)
                
                # Semantic naming for common headers without numbers
                if not current_clause_number and line.upper() in ['PURPOSE', 'SCOPE', 'RESPONSIBILITIES', 'PROCEDURE', 'RECORDS']:
                    current_clause_number = f"{document_number}.{line.upper()}"
                
                current_text_parts = []
            else:
                # Accumulate text (filter very short lines)
                if current_heading and len(line) > 20:
                    current_text_parts.append(line)
        
        # Save last section if it has content
        if current_heading and current_text_parts:
            section_text = '\n'.join(current_text_parts).strip()
            if len(section_text) >= 50:
                clauses.append({
                    "document_number": document_number,
                    "clause_number": current_clause_number if current_clause_number else f"{document_number}.{len(clauses)+1}",
                    "title": current_heading,
                    "text": section_text,
                    "characters": len(section_text)
                })
        
        return clauses
    
    def parse_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse file based on extension
        Supports: .docx, .pdf, .txt
        """
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext == 'docx':
            return self.parse_docx(file_content, filename)
        elif file_ext == 'pdf':
            return self.parse_pdf(file_content, filename)
        elif file_ext == 'txt':
            return self.parse_txt(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: docx, pdf, txt")


# Singleton instance
_qsp_parser = None

def get_qsp_parser() -> QSPParser:
    """Get singleton QSP parser instance"""
    global _qsp_parser
    if _qsp_parser is None:
        _qsp_parser = QSPParser()
    return _qsp_parser
