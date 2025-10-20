"""
QSP Compliance Checker - Real Kalibr SDK Implementation
This reimplements our QSP Compliance functionality using the actual Kalibr SDK from PyPI.

The same functionality that took 700+ lines of custom MCP code is now clean and supports
all AI models (GPT, Claude, Gemini, Copilot) automatically.
"""

from kalibr.kalibr_app import KalibrApp
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import uuid
import json
import base64
import io
from datetime import datetime, timezone
from docx import Document
import os
import re

# Initialize Kalibr App
app = KalibrApp(title="QSP Compliance Checker")

# Data Models
class QSPDocument(BaseModel):
    id: str
    filename: str
    content: str
    sections: Dict[str, str]
    upload_date: str
    processed: bool = True

class ISOSummary(BaseModel):
    id: str
    framework: str = "ISO_13485"
    version: str = "2024_summary"
    content: str
    new_clauses: List[Dict[str, str]]
    modified_clauses: List[Dict[str, str]]
    upload_date: str

class ClauseMapping(BaseModel):
    id: str
    qsp_id: str
    qsp_filename: str
    section_title: str
    iso_clause: str
    confidence_score: float
    evidence_text: str

class ComplianceGap(BaseModel):
    id: str
    qsp_id: str
    qsp_filename: str
    iso_clause: str
    gap_type: str  # "missing", "partial", "outdated"
    description: str
    severity: str  # "high", "medium", "low"
    recommendations: List[str]

# In-memory storage (replace with your database in production)
qsp_documents: Dict[str, QSPDocument] = {}
iso_summaries: Dict[str, ISOSummary] = {}
clause_mappings: Dict[str, ClauseMapping] = {}
compliance_gaps: Dict[str, ComplianceGap] = {}

# Helper functions
def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX file: {e}")

def parse_document_sections(content: str, filename: str) -> Dict[str, str]:
    """Parse document content into logical sections"""
    sections = {}
    
    # Split by common section patterns
    patterns = [
        r'\b\d+\.\s+[A-Z][^.]*\n',  # Numbered sections
        r'\b[A-Z][^.]{10,50}:?\n',   # Header patterns
        r'^[A-Z\s]{3,}\n',           # All caps headers
    ]
    
    current_section = "Introduction"
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line is a section header
        is_header = False
        for pattern in patterns:
            if re.match(pattern, line + '\n'):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = line.replace(':', '').strip()
                current_content = []
                is_header = True
                break
        
        if not is_header:
            current_content.append(line)
    
    # Save final section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    # If no sections found, treat entire document as one section
    if not sections:
        sections[f"{filename} - Full Content"] = content
    
    return sections

# Kalibr Actions - these automatically become available to all AI models

@app.action(
    name="upload_qsp_document",
    description="Upload a QSP document (.docx or .txt) for ISO 13485 compliance analysis"
)
def upload_qsp_document(filename: str, content_base64: str, file_type: str) -> dict:
    """Upload and process a QSP document"""
    try:
        # Decode base64 content
        content_bytes = base64.b64decode(content_base64)
        
        # Extract text based on file type
        if file_type.lower() == "docx":
            text_content = extract_text_from_docx(content_bytes)
        elif file_type.lower() == "txt":
            text_content = content_bytes.decode('utf-8')
        else:
            return {"success": False, "error": f"Unsupported file type: {file_type}"}
        
        # Parse into sections
        sections = parse_document_sections(text_content, filename)
        
        # Create QSP document
        qsp_doc = QSPDocument(
            id=str(uuid.uuid4()),
            filename=filename,
            content=text_content,
            sections=sections,
            upload_date=datetime.now(timezone.utc).isoformat()
        )
        
        # Store in memory
        qsp_documents[qsp_doc.id] = qsp_doc
        
        return {
            "success": True,
            "document_id": qsp_doc.id,
            "filename": filename,
            "sections_count": len(sections),
            "content_length": len(text_content),
            "message": f"QSP document '{filename}' uploaded and processed successfully"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="upload_iso_summary",
    description="Upload ISO 13485:2024 Summary of Changes for compliance benchmarking"
)
def upload_iso_summary(filename: str, content_base64: str, file_type: str) -> dict:
    """Upload ISO 13485:2024 Summary of Changes"""
    try:
        # Decode base64 content
        content_bytes = base64.b64decode(content_base64)
        
        # Extract text
        if file_type.lower() == "docx":
            text_content = extract_text_from_docx(content_bytes)
        else:
            text_content = content_bytes.decode('utf-8')
        
        # Parse ISO summary content (simplified)
        new_clauses = []
        modified_clauses = []
        
        lines = text_content.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if 'NEW CLAUSES' in line.upper() or 'NEW REQUIREMENTS' in line.upper():
                current_section = "new"
            elif 'MODIFIED CLAUSES' in line.upper() or 'CHANGED CLAUSES' in line.upper():
                current_section = "modified"
            elif line and current_section and re.match(r'\d+\.\d+', line):
                clause_info = {"clause": line, "description": line}
                if current_section == "new":
                    new_clauses.append(clause_info)
                elif current_section == "modified":
                    modified_clauses.append(clause_info)
        
        # Create ISO summary
        iso_summary = ISOSummary(
            id=str(uuid.uuid4()),
            content=text_content,
            new_clauses=new_clauses,
            modified_clauses=modified_clauses,
            upload_date=datetime.now(timezone.utc).isoformat()
        )
        
        # Store in memory
        iso_summaries[iso_summary.id] = iso_summary
        
        return {
            "success": True,
            "summary_id": iso_summary.id,
            "filename": filename,
            "new_clauses_count": len(new_clauses),
            "modified_clauses_count": len(modified_clauses),
            "message": f"ISO summary '{filename}' uploaded successfully"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="list_documents",
    description="Get all uploaded QSP documents with their processing status"
)
def list_documents(format_type: str = "summary") -> dict:
    """List all QSP documents"""
    try:
        docs = list(qsp_documents.values())
        
        if format_type == "detailed":
            return {
                "documents": [doc.dict() for doc in docs],
                "count": len(docs)
            }
        
        # Summary format
        summary = []
        for doc in docs:
            summary.append({
                "id": doc.id,
                "filename": doc.filename,
                "sections": len(doc.sections),
                "upload_date": doc.upload_date,
                "processed": doc.processed
            })
        
        return {
            "documents": summary,
            "count": len(docs),
            "message": f"Found {len(docs)} QSP documents"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="run_clause_mapping",
    description="Execute AI-powered mapping between QSP sections and ISO 13485 clauses"
)
def run_clause_mapping() -> dict:
    """Run AI-powered clause mapping analysis"""
    try:
        if not qsp_documents:
            return {"success": False, "error": "No QSP documents found. Upload documents first."}
        
        # Standard ISO 13485 clauses for mapping
        iso_clauses = [
            "4.1 General requirements",
            "4.2 Documentation requirements", 
            "5.1 Management commitment",
            "5.2 Customer focus",
            "5.3 Quality policy",
            "7.3 Design and development",
            "7.4 Purchasing",
            "8.1 General",
            "8.5 Improvement"
        ]
        
        # Clear existing mappings
        clause_mappings.clear()
        
        total_mappings = 0
        
        # Simple mapping logic (in real implementation, use AI/LLM here)
        for doc in qsp_documents.values():
            for section_title, section_content in doc.sections.items():
                if len(section_content) < 50:  # Skip very short sections
                    continue
                
                # Simple keyword-based mapping (replace with AI analysis)
                for iso_clause in iso_clauses:
                    confidence = calculate_mapping_confidence(section_content, iso_clause)
                    
                    if confidence > 0.3:  # Only store meaningful mappings
                        mapping = ClauseMapping(
                            id=str(uuid.uuid4()),
                            qsp_id=doc.id,
                            qsp_filename=doc.filename,
                            section_title=section_title,
                            iso_clause=iso_clause,
                            confidence_score=confidence,
                            evidence_text=section_content[:200] + "..."
                        )
                        
                        clause_mappings[mapping.id] = mapping
                        total_mappings += 1
        
        return {
            "success": True,
            "documents_processed": len(qsp_documents),
            "mappings_generated": total_mappings,
            "message": f"Clause mapping completed. Generated {total_mappings} mappings from {len(qsp_documents)} documents."
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

def calculate_mapping_confidence(content: str, iso_clause: str) -> float:
    """Simple confidence calculation (replace with AI analysis)"""
    content_lower = content.lower()
    clause_lower = iso_clause.lower()
    
    # Simple keyword matching
    keywords = {
        "4.1 general": ["requirements", "general", "quality", "management"],
        "4.2 documentation": ["documentation", "document", "records", "procedure"],
        "5.1 management": ["management", "commitment", "leadership", "responsibility"],
        "7.3 design": ["design", "development", "product", "specification"],
        "7.4 purchasing": ["purchasing", "supplier", "vendor", "procurement"],
        "8.1 general": ["monitoring", "measurement", "control"],
        "8.5 improvement": ["improvement", "corrective", "preventive", "action"]
    }
    
    for key, terms in keywords.items():
        if key in clause_lower:
            matches = sum(1 for term in terms if term in content_lower)
            return min(matches * 0.2, 0.9)  # Max confidence of 0.9
    
    return 0.1  # Default low confidence

@app.action(
    name="run_compliance_analysis",
    description="Analyze compliance gaps between current QSPs and ISO 13485:2024 changes"
)
def run_compliance_analysis() -> dict:
    """Run comprehensive compliance gap analysis"""
    try:
        if not iso_summaries:
            return {"success": False, "error": "No ISO summary found. Upload ISO 13485:2024 Summary first."}
        
        if not clause_mappings:
            return {"success": False, "error": "No clause mappings found. Run clause mapping first."}
        
        # Get the latest ISO summary
        latest_iso = list(iso_summaries.values())[-1]
        
        # Extract changed clauses
        changed_clauses = set()
        for clause in latest_iso.new_clauses + latest_iso.modified_clauses:
            changed_clauses.add(clause.get('clause', ''))
        
        # Find gaps
        compliance_gaps.clear()
        mapped_clauses = {mapping.iso_clause for mapping in clause_mappings.values()}
        
        gaps_found = 0
        for changed_clause in changed_clauses:
            if not changed_clause:
                continue
            
            # Check if clause is covered by mappings
            found_mapping = any(changed_clause in clause for clause in mapped_clauses)
            
            if not found_mapping:
                gap = ComplianceGap(
                    id=str(uuid.uuid4()),
                    qsp_id="",
                    qsp_filename="Multiple",
                    iso_clause=changed_clause,
                    gap_type="missing",
                    description=f"No QSP documents address the new/modified clause: {changed_clause}",
                    severity="high",
                    recommendations=[
                        f"Create or update QSP documentation to address {changed_clause}",
                        "Review existing procedures for compliance gaps",
                        "Assign responsibility for implementing new requirements"
                    ]
                )
                compliance_gaps[gap.id] = gap
                gaps_found += 1
        
        # Calculate compliance score
        total_changed = len([c for c in changed_clauses if c])
        high_conf_mappings = sum(1 for m in clause_mappings.values() if m.confidence_score > 0.7)
        overall_score = min((high_conf_mappings / max(total_changed, 1)) * 100, 100)
        
        return {
            "success": True,
            "overall_score": round(overall_score, 2),
            "gaps_found": gaps_found,
            "high_priority_gaps": gaps_found,  # All missing clauses are high priority
            "total_documents": len(qsp_documents),
            "message": f"Compliance analysis completed. Found {gaps_found} gaps with {overall_score:.1f}% overall compliance score."
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="get_compliance_status",
    description="Get current compliance overview with scores and key metrics"
)
def get_compliance_status() -> dict:
    """Get current compliance status and metrics"""
    try:
        return {
            "total_documents": len(qsp_documents),
            "total_mappings": len(clause_mappings),
            "total_gaps": len(compliance_gaps),
            "iso_summary_loaded": len(iso_summaries) > 0,
            "analysis_ready": len(qsp_documents) > 0 and len(iso_summaries) > 0,
            "last_updated": datetime.now(timezone.utc).isoformat(),
            "status": "Ready for analysis" if len(qsp_documents) > 0 and len(iso_summaries) > 0 else "Upload documents to begin"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="get_detailed_gaps",
    description="Get detailed compliance gaps with recommendations"
)
def get_detailed_gaps(severity: str = "all") -> dict:
    """Get detailed compliance gaps"""
    try:
        gaps = list(compliance_gaps.values())
        
        if severity != "all":
            gaps = [gap for gap in gaps if gap.severity == severity]
        
        if not gaps:
            return {
                "gaps": [],
                "count": 0,
                "message": "No compliance gaps found. Excellent compliance status!"
            }
        
        return {
            "gaps": [gap.dict() for gap in gaps[:10]],  # Limit to first 10
            "count": len(gaps),
            "message": f"Found {len(gaps)} compliance gaps"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}

@app.action(
    name="query_specific_clause",
    description="Query compliance status for a specific ISO clause"
)
def query_specific_clause(clause: str, include_recommendations: bool = True) -> dict:
    """Query compliance for a specific ISO clause"""
    try:
        # Find mappings for this clause
        related_mappings = [
            mapping for mapping in clause_mappings.values()
            if clause.lower() in mapping.iso_clause.lower()
        ]
        
        # Find gaps for this clause
        related_gaps = [
            gap for gap in compliance_gaps.values()
            if clause.lower() in gap.iso_clause.lower()
        ]
        
        result = {
            "clause": clause,
            "mappings_found": len(related_mappings),
            "gaps_found": len(related_gaps),
            "compliance_status": "compliant" if related_mappings and not related_gaps else "non-compliant"
        }
        
        if related_mappings:
            result["mapped_documents"] = [
                {
                    "filename": mapping.qsp_filename,
                    "section": mapping.section_title,
                    "confidence": mapping.confidence_score
                }
                for mapping in related_mappings[:3]  # Top 3
            ]
        
        if related_gaps and include_recommendations:
            result["gaps"] = [
                {
                    "description": gap.description,
                    "severity": gap.severity,
                    "recommendations": gap.recommendations
                }
                for gap in related_gaps[:2]  # Top 2
            ]
        
        return {"success": True, **result}
        
    except Exception as e:
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    print("QSP Compliance Checker - Kalibr Implementation")
    print("Run with: python -m kalibr serve qsp_compliance_kalibr.py")
    print("\nThis will automatically provide:")
    print("✅ Claude MCP WebSocket endpoint")
    print("✅ GPT Actions OpenAPI schema")  
    print("✅ Gemini Extensions integration")
    print("✅ Microsoft Copilot plugin support")
    print("\nAll from this single file with clean function definitions!")