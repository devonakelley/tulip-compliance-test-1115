"""
Document processing engine for Enterprise QSP Compliance System
Handles document upload, parsing, and content extraction
"""

import hashlib
import mimetypes
import asyncio
import logging
from typing import List, Dict, Optional, Any, Union
from datetime import datetime, timezone
from pathlib import Path
import uuid
import json
import re

# Document processing libraries
from docx import Document as DocxDocument
import PyPDF2
import pandas as pd
from bs4 import BeautifulSoup

# Database imports
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, delete
from database.models import Document, DocumentSection, DocumentTypeEnum
from models import DocumentMetadata, DocumentType, UploadResponse
from config import settings
from utils import FileUtils, TextProcessor
from rag import RAGEngine

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles all document processing operations"""
    
    def __init__(self):
        self.supported_types = {
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.xlsx': self._process_excel,
            '.csv': self._process_csv,
            '.html': self._process_html,
            '.xml': self._process_xml
        }
        self.file_utils = FileUtils()
        self.text_processor = TextProcessor()
        self.rag_engine = RAGEngine()
        
    async def process_upload(
        self,
        file_data: bytes,
        filename: str,
        document_type: str,
        user_id: str,
        session: AsyncSession,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Process uploaded document
        
        Args:
            file_data: Raw file bytes
            filename: Original filename
            document_type: Type of document
            user_id: ID of uploading user
            session: Database session
            metadata: Optional metadata
            
        Returns:
            Document ID
        """
        try:
            # Validate file
            await self._validate_file(file_data, filename)
            
            # Generate file hash
            file_hash = hashlib.sha256(file_data).hexdigest()
            
            # Check for duplicates
            duplicate = await self._check_duplicate(file_hash, user_id, session)
            if duplicate:
                logger.info(f"Duplicate file detected: {filename}")
                return duplicate.id
            
            # Save file to storage
            file_path = await self._save_file(file_data, filename, user_id)
            
            # Create document record
            document = Document(
                id=uuid.uuid4(),
                filename=filename,
                document_type=DocumentTypeEnum(document_type),
                file_path=str(file_path),
                file_size=len(file_data),
                file_hash=file_hash,
                user_id=uuid.UUID(user_id),
                processing_status="uploaded",
                metadata=metadata or {}
            )
            
            session.add(document)
            await session.commit()
            
            logger.info(f"Document uploaded successfully: {document.id}")
            return str(document.id)
            
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            await session.rollback()
            raise
    
    async def process_document_async(
        self,
        document_id: str,
        session: AsyncSession
    ) -> bool:
        """
        Asynchronously process document content
        
        Args:
            document_id: Document ID to process
            session: Database session
            
        Returns:
            Success status
        """
        try:
            # Get document
            document = await self._get_document(document_id, session)
            if not document:
                raise ValueError(f"Document not found: {document_id}")
            
            # Update status
            await self._update_status(document_id, "processing", session)
            
            # Extract content
            content = await self._extract_content(document)
            
            # Parse sections
            sections = await self._parse_sections(content, document)
            
            # Update document
            await self._update_document_content(
                document_id, content, sections, session
            )
            
            # Update status
            await self._update_status(document_id, "completed", session)
            
            logger.info(f"Document processed successfully: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            await self._update_status(document_id, "failed", session, str(e))
            return False
    
    async def list_documents(
        self,
        user_id: str,
        session: AsyncSession,
        document_type: Optional[str] = None,
        limit: int = 50,
        offset: int = 0
    ) -> List[DocumentMetadata]:
        """
        List user documents with filtering
        
        Args:
            user_id: User ID
            document_type: Optional document type filter
            limit: Maximum number of results
            offset: Offset for pagination
            session: Database session
            
        Returns:
            List of document metadata
        """
        try:
            query = select(Document).where(
                Document.user_id == uuid.UUID(user_id),
                Document.is_active
            )
            
            if document_type:
                query = query.where(Document.document_type == DocumentTypeEnum(document_type))
            
            query = query.order_by(Document.upload_date.desc()).offset(offset).limit(limit)
            
            result = await session.execute(query)
            documents = result.scalars().all()
            
            return [
                DocumentMetadata(
                    document_id=str(doc.id),
                    filename=doc.filename,
                    document_type=DocumentType(doc.document_type.value),
                    file_size=doc.file_size,
                    upload_date=doc.upload_date,
                    processed_date=doc.processed_date,
                    user_id=str(doc.user_id),
                    status=doc.processing_status,
                    sections_count=doc.sections_count,
                    clause_mappings_count=doc.clause_mappings_count,
                    compliance_score=doc.compliance_score,
                    last_analysis_date=doc.last_analysis_date,
                    tags=doc.tags or [],
                    metadata=doc.metadata or {}
                )
                for doc in documents
            ]
            
        except Exception as e:
            logger.error(f"Failed to list documents: {e}")
            return []
    
    async def delete_document(
        self,
        document_id: str,
        user_id: str,
        session: AsyncSession
    ) -> bool:
        """
        Delete document and associated data
        
        Args:
            document_id: Document ID
            user_id: User ID (for authorization)
            session: Database session
            
        Returns:
            Success status
        """
        try:
            # Get document
            document = await self._get_document(document_id, session)
            if not document or str(document.user_id) != user_id:
                return False
            
            # Soft delete (mark as inactive)
            await session.execute(
                update(Document)
                .where(Document.id == uuid.UUID(document_id))
                .values(is_active=False)
            )
            
            # Remove file from storage
            if document.file_path:
                await self._delete_file(document.file_path)
            
            await session.commit()
            
            logger.info(f"Document deleted: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False
    
    async def get_document_content(
        self,
        document_id: str,
        user_id: str,
        session: AsyncSession
    ) -> Optional[Dict[str, Any]]:
        """
        Get document content and sections
        
        Args:
            document_id: Document ID
            user_id: User ID (for authorization)
            session: Database session
            
        Returns:
            Document content and sections
        """
        try:
            document = await self._get_document(document_id, session)
            if not document or str(document.user_id) != user_id:
                return None
            
            # Get sections
            sections_query = select(DocumentSection).where(
                DocumentSection.document_id == uuid.UUID(document_id)
            ).order_by(DocumentSection.section_order)
            
            sections_result = await session.execute(sections_query)
            sections = sections_result.scalars().all()
            
            return {
                "document": {
                    "id": str(document.id),
                    "filename": document.filename,
                    "document_type": document.document_type.value,
                    "content": document.content,
                    "processed_content": document.processed_content,
                    "upload_date": document.upload_date,
                    "processed_date": document.processed_date
                },
                "sections": [
                    {
                        "id": str(section.id),
                        "section_number": section.section_number,
                        "title": section.title,
                        "content": section.content,
                        "level": section.level,
                        "word_count": section.word_count,
                        "confidence_score": section.confidence_score
                    }
                    for section in sections
                ]
            }
            
        except Exception as e:
            logger.error(f"Failed to get document content: {e}")
            return None
    
    # Private methods
    
    async def _validate_file(self, file_data: bytes, filename: str):
        """Validate uploaded file"""
        # Check file size
        if len(file_data) > settings.MAX_FILE_SIZE_MB * 1024 * 1024:
            raise ValueError(f"File too large. Maximum size: {settings.MAX_FILE_SIZE_MB}MB")
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in settings.ALLOWED_FILE_TYPES:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Check if file is empty
        if len(file_data) == 0:
            raise ValueError("Empty file")
    
    async def _check_duplicate(
        self, 
        file_hash: str, 
        user_id: str, 
        session: AsyncSession
    ) -> Optional[Document]:
        """Check for duplicate files"""
        query = select(Document).where(
            Document.file_hash == file_hash,
            Document.user_id == uuid.UUID(user_id),
            Document.is_active
        )
        
        result = await session.execute(query)
        return result.scalar_one_or_none()
    
    async def _save_file(self, file_data: bytes, filename: str, user_id: str) -> Path:
        """Save file to storage"""
        # Create user directory
        user_dir = Path(settings.UPLOAD_DIRECTORY) / user_id
        user_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_ext = Path(filename).suffix
        unique_filename = f"{timestamp}_{uuid.uuid4().hex[:8]}{file_ext}"
        
        file_path = user_dir / unique_filename
        
        # Write file
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        return file_path
    
    async def _delete_file(self, file_path: str):
        """Delete file from storage"""
        try:
            Path(file_path).unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Failed to delete file {file_path}: {e}")
    
    async def _get_document(self, document_id: str, session: AsyncSession) -> Optional[Document]:
        """Get document by ID"""
        query = select(Document).where(Document.id == uuid.UUID(document_id))
        result = await session.execute(query)
        return result.scalar_one_or_none()
    
    async def _update_status(
        self, 
        document_id: str, 
        status: str, 
        session: AsyncSession,
        error_message: Optional[str] = None
    ):
        """Update document processing status"""
        values = {"processing_status": status}
        
        if status == "completed":
            values["processed_date"] = datetime.now(timezone.utc)
        
        if error_message:
            values["processing_error"] = error_message
        
        await session.execute(
            update(Document)
            .where(Document.id == uuid.UUID(document_id))
            .values(**values)
        )
        await session.commit()
    
    async def _extract_content(self, document: Document) -> str:
        """Extract text content from document"""
        file_ext = Path(document.filename).suffix.lower()
        
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        processor = self.supported_types[file_ext]
        return await processor(document.file_path)
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            logger.error(f"Failed to process DOCX file: {e}")
            raise
    
    async def _process_text(self, file_path: str) -> str:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file")
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF file"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text.strip())
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Failed to process PDF file: {e}")
            raise
    
    async def _process_excel(self, file_path: str) -> str:
        """Process Excel file"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            content_parts = []
            
            for sheet_name, df in excel_data.items():
                content_parts.append(f"Sheet: {sheet_name}")
                content_parts.append(df.to_string(index=False))
                content_parts.append("")
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"Failed to process Excel file: {e}")
            raise
    
    async def _process_csv(self, file_path: str) -> str:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Failed to process CSV file: {e}")
            raise
    
    async def _process_html(self, file_path: str) -> str:
        """Process HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            return soup.get_text(separator='\n', strip=True)
            
        except Exception as e:
            logger.error(f"Failed to process HTML file: {e}")
            raise
    
    async def _process_xml(self, file_path: str) -> str:
        """Process XML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'xml')
            return soup.get_text(separator='\n', strip=True)
            
        except Exception as e:
            logger.error(f"Failed to process XML file: {e}")
            raise
    
    async def _parse_sections(self, content: str, document: Document) -> List[Dict[str, Any]]:
        """Parse document content into sections"""
        sections = []
        
        # Use text processor to identify sections
        parsed_sections = await self.text_processor.parse_sections(
            content, document.document_type.value
        )
        
        for i, section_data in enumerate(parsed_sections):
            section = {
                "section_number": section_data.get("number"),
                "title": section_data.get("title", f"Section {i+1}"),
                "content": section_data.get("content", ""),
                "level": section_data.get("level", 1),
                "word_count": len(section_data.get("content", "").split()),
                "confidence_score": section_data.get("confidence", 0.8)
            }
            sections.append(section)
        
        return sections
    
    async def _update_document_content(
        self,
        document_id: str,
        content: str,
        sections: List[Dict[str, Any]],
        session: AsyncSession
    ):
        """Update document with processed content and sections"""
        try:
            # Update document
            await session.execute(
                update(Document)
                .where(Document.id == uuid.UUID(document_id))
                .values(
                    content=content,
                    processed_content={"sections_count": len(sections)},
                    sections_count=len(sections),
                    processed_date=datetime.now(timezone.utc)
                )
            )
            
            # Create sections
            for i, section_data in enumerate(sections):
                section = DocumentSection(
                    id=uuid.uuid4(),
                    document_id=uuid.UUID(document_id),
                    section_number=section_data["section_number"],
                    title=section_data["title"],
                    content=section_data["content"],
                    content_hash=hashlib.sha256(section_data["content"].encode()).hexdigest(),
                    level=section_data["level"],
                    section_order=i,
                    word_count=section_data["word_count"],
                    confidence_score=section_data["confidence_score"]
                )
                session.add(section)
            
            await session.commit()
            
        except Exception as e:
            logger.error(f"Failed to update document content: {e}")
            raise
    
    async def process_document_with_rag(
        self,
        document_id: str,
        session
    ) -> Dict[str, Any]:
        """
        Process document with RAG system for semantic search
        
        Args:
            document_id: Document ID to process
            session: Database session
            
        Returns:
            Processing results
        """
        try:
            # Get document from database
            document = await session.documents.find_one({"_id": document_id})
            
            if not document:
                raise ValueError(f"Document not found: {document_id}")
            
            if document.get("processing_status") != "completed":
                raise ValueError(f"Document not ready for RAG processing: {document_id}")
            
            logger.info(f"Processing document with RAG: {document_id}")
            
            # Process with RAG engine
            rag_result = await self.rag_engine.process_qsp_document(document)
            
            if rag_result.get('success'):
                # Update document status
                await session.documents.update_one(
                    {"_id": document_id},
                    {
                        "$set": {
                            "rag_processed": True,
                            "rag_processing_date": datetime.now(timezone.utc),
                            "rag_chunks_count": rag_result.get('chunks_created', 0),
                            "rag_sections_count": rag_result.get('sections_identified', 0)
                        }
                    }
                )
                
                logger.info(f"Document RAG processing completed: {document_id}")
            
            return rag_result
            
        except Exception as e:
            logger.error(f"Failed to process document with RAG: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_id': document_id
            }
